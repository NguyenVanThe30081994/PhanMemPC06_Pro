# -*- coding: utf-8 -*-
"""
outline_parser.py — Bộ phân tích cấu trúc đề cương (.docx / .txt)

Thuật toán tự động nhận diện cấu trúc phân cấp bất kể số tầng:
    - Cấp 1: I., II., III. (La Mã)
    - Cấp 2: 1., 2., 3. (số)
    - Cấp 3: 1.1., 1.2., 2.1. (số kép)
    - Cấp 4: 1.1.1., 2.2.1. (số ba) ... và cứ tiếp tục
    - Gạch đầu dòng: "-", "–", "•"
    - Mục con (+): "+"  -> thuộc gạch đầu dòng cha gần nhất
    - Đoạn văn tự do: nội dung thường thuộc mục đang mở

Trả về cây JSON:
{
  "title": ...,
  "subtitle": ...,
  "sections": [
    {
      "id": "...",
      "type": "h1|h2|h3|h4|bullet|plus|para",
      "label": "I" | "1" | "1.1" | "2.2.1",
      "text": "nội dung",
      "children": [...]
    }
  ]
}
"""

import io
import re
import uuid

try:
    from docx import Document as _DocxDocument
except ImportError:
    _DocxDocument = None


# ── Nhận diện cấp heading ────────────────────────────────────────────────
# Số La Mã cho cấp 1 (I., II., III., IV., ...)
_RE_ROMAN = re.compile(r'^([IVXLCDM]+)\.?\s+(.+)$')

# Số thập phân: 1.1.1., 1.1., 1. (đếm số chấm -> cấp độ)
_RE_DECIMAL = re.compile(r'^((?:\d{1,3}\.){1,6}\d{1,3})\.?\s+(.+)$')

# Số đơn giản: 1., 2. (dùng khi không khớp decimal)
_RE_SIMPLE_NUM = re.compile(r'^(\d{1,3})\.\s+(.+)$')

# Dấu cộng / gạch đầu dòng
_RE_PLUS = re.compile(r'^\+[\s:：]*\s*(.*)$')
_RE_DASH = re.compile(r'^[-–—•*]\s+(.*)$')

# Ký tự La Mã hợp lệ (để tránh nhận nhầm chữ thường)
_ROMAN_TOKENS = {'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV', 'XVI'}


def _roman_value(token):
    """Chuyển La Mã -> số, trả None nếu không hợp lệ."""
    token = token.upper()
    values = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
    total = 0
    prev = 0
    for ch in reversed(token):
        v = values.get(ch)
        if v is None:
            return None
        if v < prev:
            total -= v
        else:
            total += v
        prev = v
    if total < 1 or total > 3999:
        return None
    return total


def classify_line(text):
    """
    Phân loại một dòng văn bản.
    Trả về (type, label, content):
        type: 'h1'|'h2'|'h3'|'h4'|'h5'|'h6'|'bullet'|'plus'|'para'
    """
    raw = str(text or '').strip()
    if not raw:
        return ('skip', None, '')

    # 1. Số La Mã -> cấp 1
    m = _RE_ROMAN.match(raw)
    if m and _roman_value(m.group(1)) is not None:
        return ('h1', m.group(1).upper(), m.group(2).strip())

    # 2. Số thập phân 1.1.1., 1.1., ...
    #    "1."   -> mục cấp 2 (h2), "1.1" -> cấp 3 (h3), "1.1.1" -> cấp 4 (h4)
    m = _RE_DECIMAL.match(raw)
    if m:
        num = m.group(1)
        depth = num.count('.')           # số chấm: "1.1" -> 1, "1.1.1" -> 2
        lvl = min(depth + 2, 9)          # "1.1" -> h3, "1.1.1" -> h4
        return ('h%d' % lvl, num, m.group(2).strip())

    # 3. Số đơn 1., 2. -> cấp 2 (nếu chưa khớp ở trên)
    m = _RE_SIMPLE_NUM.match(raw)
    if m:
        return ('h2', m.group(1), m.group(2).strip())

    # 4. Gạch đầu dòng
    m = _RE_DASH.match(raw)
    if m:
        return ('bullet', '–', m.group(1).strip())

    # 5. Dấu cộng (mục con thuộc gạch đầu dòng cha)
    m = _RE_PLUS.match(raw)
    if m:
        content = m.group(1).strip()
        if content:
            return ('plus', '+', content)
        return ('plus', '+', '')

    # 6. Đoạn văn tự do
    return ('para', None, raw)


_HEADING_DEPTH = {
    'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6, 'h7': 7, 'h8': 8, 'h9': 9,
}


def _is_heading(kind):
    return kind in _HEADING_DEPTH


def _node(kind, label, text):
    return {
        'id': uuid.uuid4().hex[:8],
        'type': kind,
        'label': label or '',
        'text': text or '',
        'children': [],
    }


def build_tree(paragraphs):
    """
    paragraphs: iterable các dòng văn bản (str) đã strip.
    Trả về dict {title, subtitle, sections, stats}.
    """
    root = _node('root', '', '')
    title_lines = []
    stack = [root]                # ngăn xếp các heading đang mở
    last_bullet_at = {}           # id(heading) -> nút bullet gần nhất (để gắn +)

    for text in paragraphs:
        text = str(text or '').strip()
        if not text:
            continue

        kind, label, content = classify_line(text)
        if kind == 'skip':
            continue

        # ── Heading: đẩy vào đúng vị trí cây ──
        if _is_heading(kind):
            depth = _HEADING_DEPTH[kind]
            # Pop các heading có độ sâu >= depth
            while len(stack) > 1 and _HEADING_DEPTH.get(stack[-1]['type'], 99) >= depth:
                stack.pop()
            node = _node(kind, label, content)
            stack[-1]['children'].append(node)
            stack.append(node)
            continue

        # ── Nội dung trước heading đầu tiên -> title block ──
        has_heading = any(_is_heading(n['type']) for n in stack[1:])
        if not has_heading:
            title_lines.append(text)
            continue

        parent = stack[-1]

        # ── Dấu + : gắn vào gạch đầu dòng cha gần nhất ──
        if kind == 'plus':
            lb = last_bullet_at.get(id(parent))
            if lb is not None:
                lb['children'].append(_node(kind, label, content))
            else:
                parent['children'].append(_node(kind, label, content))
            continue

        # ── Gạch đầu dòng / đoạn văn: thuộc heading đang mở ──
        node = _node(kind, label, content)
        parent['children'].append(node)
        if kind == 'bullet':
            last_bullet_at[id(parent)] = node
        elif kind == 'para':
            # Đoạn văn không phải bullet -> reset bullet gần nhất
            last_bullet_at[id(parent)] = None

    result = {
        'title': title_lines[0] if title_lines else '(Chưa có tiêu đề)',
        'subtitle': title_lines[1] if len(title_lines) > 1 else '',
        'sections': root['children'],
    }
    _normalize_levels(result['sections'])
    return result


def _normalize_levels(nodes):
    """
    Nếu cây không có cấp h1 (ví dụ đề cương bắt đầu bằng '1.', '2.'),
    tự động shift toàn bộ cấp độ lên một bậc để luôn có h1 làm gốc.
    Chỉ thực hiện khi cấp nhỏ nhất hiện có > h1.
    """
    if not nodes:
        return
    min_level = 99
    max_level = 0

    def scan(ns):
        nonlocal min_level, max_level
        for n in ns:
            if n['type'] in _HEADING_DEPTH:
                lv = _HEADING_DEPTH[n['type']]
                min_level = min(min_level, lv)
                max_level = max(max_level, lv)
            scan(n.get('children') or [])

    scan(nodes)
    if min_level == 99 or min_level <= 1:
        return  # đã có h1 hoặc không có heading

    shift = min_level - 1  # số bậc cần shift lên

    def lift(ns):
        for n in ns:
            if n['type'] in _HEADING_DEPTH:
                new_lv = min(_HEADING_DEPTH[n['type']] - shift, 9)
                n['type'] = 'h%d' % new_lv
            lift(n.get('children') or [])

    lift(nodes)


def parse_docx(filepath_or_bytes):
    """Đọc .docx -> cây cấu trúc."""
    if _DocxDocument is None:
        raise RuntimeError('Chưa cài thư viện python-docx. Hãy cài: pip install python-docx')

    if isinstance(filepath_or_bytes, (bytes, bytearray)):
        document = _DocxDocument(io.BytesIO(bytes(filepath_or_bytes)))
    else:
        document = _DocxDocument(filepath_or_bytes)

    paragraphs = []
    for para in document.paragraphs:
        text = str(getattr(para, 'text', '') or '').strip()
        if text:
            paragraphs.append(text)

    return build_tree(paragraphs)


def parse_text(raw_text):
    """Đọc văn bản thuần (.txt) -> cây cấu trúc."""
    lines = [line.strip() for line in str(raw_text or '').splitlines()]
    return build_tree(lines)


# ── Tiện ích thống kê ────────────────────────────────────────────────────
def collect_stats(tree):
    counts = {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0, 'bullet': 0, 'plus': 0, 'para': 0}

    def walk(nodes):
        for n in nodes:
            if n['type'] in counts:
                counts[n['type']] += 1
            walk(n.get('children') or [])

    walk(tree.get('sections') or [])
    return counts


def dump_tree(tree, max_len=80):
    """In cây ra console để kiểm tra (debug)."""
    lines = []
    lines.append('TITLE: ' + (tree.get('title') or ''))
    if tree.get('subtitle'):
        lines.append('SUB: ' + tree['subtitle'])

    def walk(nodes, ind=0):
        for n in nodes:
            mark = {'h1': '█', 'h2': '▸', 'h3': '·', 'h4': '˚',
                    'h5': '•', 'h6': '◦', 'bullet': '-', 'plus': '+',
                    'para': '¶'}.get(n['type'], '?')
            lbl = (n['label'] + ' ') if n.get('label') else ''
            txt = (n.get('text') or '').replace('\n', ' ')[:max_len]
            lines.append('  ' * ind + f'{mark} {lbl}{txt}')
            walk(n.get('children') or [], ind + 1)

    walk(tree.get('sections') or [])
    return '\n'.join(lines)


if __name__ == '__main__':
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else None
    if not path:
        print('Cách dùng: python outline_parser.py <file.docx|file.txt>')
        sys.exit(1)
    if path.lower().endswith('.docx'):
        result = parse_docx(path)
    else:
        with open(path, encoding='utf-8') as f:
            result = parse_text(f.read())
    print(dump_tree(result))
    print('\nSTATS:', collect_stats(result))
