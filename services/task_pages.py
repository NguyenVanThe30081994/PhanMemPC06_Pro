# -*- coding: utf-8 -*-
"""
Cụm page handler task: trang danh sách/chỉnh sửa công việc (_tasks_page_v2), chi tiết
task (_task_detail_v2), tạo đầu mục đề cương (_create_outline_items_v2), xem trước
import đề cương (_preview_outline_import_v2), cập nhật trạng thái (_update_task_status_v2),
nộp báo cáo (_submit_task_report_v2), xuất biểu mẫu/Word (_export_form_task_v2,
_export_outline_word_v2), ma trận tiến độ đề cương (_build_outline_progress_matrix).

Tách từ routes/tasks.py (Pha 2). routes/tasks.py re-export các tên còn dùng.
"""

import io
import json
from datetime import datetime

from flask import flash, redirect, request, send_file, session, url_for
from sqlalchemy.orm import joinedload
from werkzeug.utils import secure_filename
try:
    from openpyxl import Workbook
except ImportError:
    Workbook = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from models import (
    AppRole,
    Task,
    TaskAssignment,
    TaskComment,
    TaskFormField,
    TaskItem,
    TaskSubmission,
    TaskSubmissionFile,
    User,
    db,
)
from category_helpers import (
    apply_reference_display,
    canonicalize_category_value,
    module_category_options,
    stable_form_category_options,
    sync_record_categories,
)
from google_forms import extract_google_form_id
from permissions import current_is_admin
from report_cycles import config_to_json as report_config_to_json
from task_blueprints import (
    workflow_blueprint_example_catalog,
    workflow_blueprint_form_field_defs,
    workflow_blueprint_item_configs,
    workflow_blueprint_report_schema,
    workflow_blueprint_summary_text,
    workflow_blueprint_task_mode,
)
from task_page_builders import (
    build_task_detail_page_context,
    build_task_list_page_context,
    prepare_task_workspace_record,
    task_visible_for_user,
)
from task_workspace import (
    build_task_workspace_attrs,
)
from utils import (
    push_notif,
    remove_accents,
    render_auto_template as render_template,
)
from routes.email_service import send_task_assignment_emails

logger = __import__('logging').getLogger(__name__)

from services.blueprint_parsing import _parse_task_workflow_blueprint_from_request
from services.outline_engine import (
    _outline_blank_numeric,
    _outline_sources_json,
    _parse_outline_blank_value,
    _parse_outline_upload_titles,
)
from services.outline_rows import _parse_outline_upload_rows
from services.outline_submission import (
    _find_report_secondary_linked_item,
    _outline_merged_content,
    _outline_submission_values,
    _propagate_submission_to_linked_items,
)
from services.task_assignees import (
    _create_assignment_records,
    _resolve_assignees,
    _resolve_managers,
    _resolve_viewers,
)
from services.task_categories import (
    _decorate_task_categories,
    _task_domain_options,
    _task_field_options,
    _task_priority_options,
    _task_type_options,
)
from services.task_deadline import (
    _computed_task_deadline,
    _parse_deadline,
    _parse_task_report_period_from_request,
    _task_current_cycle,
    _task_report_kind_label,
    _task_report_period,
)
from services.task_form_fields import (
    _normalize_task_form_field_type,
    _task_form_fields_for_user,
)
from services.task_google_forms import (
    _apply_task_google_form_view_state,
    _hydrate_google_form_fields,
    _normalize_google_form_match_mode,
    _parse_google_form_builder_schema,
    _task_form_field_db_kwargs,
)
from services.task_guards import (
    _can_delete_task,
    _can_edit_task,
    _can_manage_task,
    _can_watch_task,
)
from services.task_import_draft_helpers import _json_dump
from services.task_import_drafts import _validate_task_visibility_before_publish
from services.task_modes import (
    TASK_ASSIGNMENT_STATUS_LABELS,
    TASK_MODE_DEFAULT,
    _requested_task_mode,
    _task_assignment_display_status,
    _task_assignment_status_class,
    _task_assignment_status_label,
    _task_mode,
    _task_mode_description,
    _task_mode_label,
)
from services.task_permissions import (
    _can_process_task_module,
    _can_view_all_tasks,
    _current_perms,
)
from services.task_report_schema import (
    CHILD_TASK_ALLOWED_REPORT_KINDS,
    _load_task_report_schema,
    _parse_task_report_schema_from_request,
    _report_checkbox_value,
)
from services.task_report_views import (
    _build_assignment_report_context,
    _build_structured_task_report_comment,
    _build_structured_task_report_form,
    _parse_structured_file_report_submission,
)
from services.task_runtime_sync import _lazy_repair_task_runtime
from services.task_scope import (
    _parse_bulk_child_task_titles,
    _requested_manager_role_ids,
    _requested_manager_user_ids,
    _requested_role_ids,
    _requested_user_ids,
    _requested_viewer_role_ids,
    _requested_viewer_user_ids,
    _store_assignment_scope,
    _store_manager_scope,
    _store_viewer_scope,
)
from services.task_units import _task_assignee_unit_name
from services.task_workspace_helpers import (
    _build_rebuilt_task_summary,
    _filter_assignment_rows_for_executor_scope,
    _filter_outline_groups_for_executor_scope,
    _store_uploaded_task_file,
    _sync_assignment_group_submission,
    _task_assignment_progress_groups,
    _task_delivery_contract_groups,
    _task_is_submitted,
)
from services.task_workspace_views import (
    _build_file_task_rows,
    _build_form_task_rows,
    _build_outline_group_rows,
    _clear_outline_import_preview,
    _get_outline_import_preview,
    _outline_item_number_fields,
    _outline_item_table_cells,
    _outline_table_schema_map,
    _parse_outline_item_configs_from_request,
    _parse_outline_item_rows,
    _parse_task_form_fields_from_request,
    _resolve_outline_item_assignment,
    _set_outline_import_preview,
    _task_detail_context,
    _task_form_field_views,
    _task_form_field_views_for_user,
    _task_form_value_is_empty,
    _task_item_synthesis_text,
)

def _tasks_page_v2():
    perms = _current_perms()
    can_view_all_tasks = _can_view_all_tasks(perms)
    is_lead = _can_process_task_module(perms)
    is_admin = bool(current_is_admin())
    current_user = db.session.get(User, session["uid"])

    task_fields = _task_field_options()
    pro_units = _task_domain_options()
    task_types = _task_type_options()
    priority_items = _task_priority_options()
    active_users = []
    roles = []

    if is_lead or is_admin:
        active_users = User.query.filter_by(is_active=True).order_by(User.unit_area.asc(), User.fullname.asc()).all()
        active_users = apply_reference_display(
            sync_record_categories(
                active_users,
                module_category_options("contacts", "unit_name", "Đơn vị"),
                attr_name="unit_area",
                prefer_stable=True,
            ),
            "unit_area",
            module_category_options("contacts", "unit_name", "Đơn vị"),
            display_attr="unit_area_display",
            fallback_label="Chưa có đơn vị",
        )
        roles = AppRole.query.order_by(AppRole.name.asc()).all()

    if request.method == "POST" and is_lead:
        title = (request.form.get("title") or "").strip()
        task_mode = _requested_task_mode(request.form)
        form_provider = str(request.form.get("form_provider") or "internal").strip().lower()
        category = canonicalize_category_value(request.form.get("category") or "", task_fields, prefer_stable=True)
        domain = canonicalize_category_value(
            request.form.get("unit_name") or request.form.get("domain") or "",
            pro_units,
            prefer_stable=True,
        )
        content = (request.form.get("description") or request.form.get("content") or "").strip()
        priority = canonicalize_category_value(request.form.get("priority") or "Trung bình", priority_items, prefer_stable=True)
        task_type = canonicalize_category_value(request.form.get("task_type") or "Công việc thường xuyên", task_types, prefer_stable=True)
        try:
            workflow_blueprint = _parse_task_workflow_blueprint_from_request(request.form)
        except ValueError as blueprint_error:
            flash(str(blueprint_error), "danger")
            return redirect(url_for("tasks_bp.tasks"))

        if workflow_blueprint:
            task_mode = workflow_blueprint_task_mode(workflow_blueprint)
            title = title or workflow_blueprint.get("title", "")
            if not content:
                content = workflow_blueprint_summary_text(workflow_blueprint)

        if not title:
            flash("Tiêu đề công việc không được để trống.", "danger")
            return redirect(url_for("tasks_bp.tasks"))

        try:
            report_schema = _parse_task_report_schema_from_request(request.form)
        except ValueError as report_schema_error:
            flash(str(report_schema_error), "danger")
            return redirect(url_for("tasks_bp.tasks"))
        if not report_schema and workflow_blueprint:
            report_schema = workflow_blueprint_report_schema(workflow_blueprint)

        google_form_builder = None
        google_form_field_defs = []
        google_form_url = ""
        google_form_id = ""
        google_form_match_mode = "unit"
        google_form_match_field = ""
        if task_mode == "FORM" and form_provider == "google":
            google_form_url = str(request.form.get("google_form_url") or "").strip()[:500]
            try:
                google_form_builder = _parse_google_form_builder_schema(
                    request.form.get("google_form_builder_json"),
                    fallback_title=title,
                    fallback_description=content,
                )
            except ValueError as builder_error:
                flash(str(builder_error), "danger")
                return redirect(url_for("tasks_bp.tasks"))
            google_form_field_defs = _hydrate_google_form_fields(google_form_builder)
            google_form_id = extract_google_form_id(google_form_url)
            builder_matching = google_form_builder.get("matching") if isinstance(google_form_builder.get("matching"), dict) else {}
            google_form_match_mode = _normalize_google_form_match_mode(
                request.form.get("google_form_match_mode") or builder_matching.get("mode") or "unit"
            )
            google_form_match_field = str(
                request.form.get("google_form_match_field")
                or builder_matching.get("match_field")
                or ""
            ).strip()[:255]

        managers, manager_error_message = _resolve_managers(request.form)
        if manager_error_message:
            flash(manager_error_message, "danger")
            return redirect(url_for("tasks_bp.tasks"))
        viewers, viewer_error_message = _resolve_viewers(request.form)
        if viewer_error_message:
            flash(viewer_error_message, "danger")
            return redirect(url_for("tasks_bp.tasks"))

        assignees = []
        blueprint_items = workflow_blueprint_item_configs(workflow_blueprint) if workflow_blueprint else []
        assign_type = request.form.get("assign_type", "unit")
        assign_role_ids = _requested_role_ids(request.form)
        assign_user_ids = _requested_user_ids(request.form)
        if task_mode in {"FILE", "FORM"} or blueprint_items:
            assignees, error_message = _resolve_assignees(request.form, domain)
            if error_message:
                flash(error_message, "danger")
                return redirect(url_for("tasks_bp.tasks"))

        attachment = request.files.get("task_file") or request.files.get("file")
        attachment_name = ""
        if attachment and attachment.filename:
            attachment_meta = _store_uploaded_task_file(attachment, "task", "template", prefix="task")
            attachment_name = attachment_meta["stored_name"] if attachment_meta else ""

        new_task = Task(
            category=category,
            domain=domain,
            title=title,
            content=content,
            deadline=_computed_task_deadline(request.form, task_type=task_type) or _parse_deadline(request.form),
            file_path=attachment_name,
            author_id=session["uid"],
            author_name=session.get("fullname", "Quản trị"),
            priority=priority,
            task_type=task_type,
            initial_status="Chưa tiếp nhận",
            task_mode=task_mode,
            form_provider=form_provider if task_mode == "FORM" else "internal",
        )
        report_period = _parse_task_report_period_from_request(request.form, task_type=task_type)
        if report_period:
            new_task.report_period_json = report_config_to_json(report_period)
        if report_schema:
            new_task.report_schema_json = json.dumps(report_schema, ensure_ascii=False)
        if task_mode == "FORM" and form_provider == "google":
            new_task.google_form_url = google_form_url or None
            new_task.google_form_id = google_form_id or None
            new_task.google_form_match_mode = google_form_match_mode
            new_task.google_form_match_field = google_form_match_field or None
            new_task.google_form_builder_json = _json_dump(google_form_builder)
        _store_assignment_scope(
            new_task,
            request.form.get("assign_type", "unit"),
            domain=domain,
            role_ids=_requested_role_ids(request.form),
            user_ids=_requested_user_ids(request.form),
        )
        _store_viewer_scope(
            new_task,
            request.form.get("viewer_scope_mode", "none"),
            role_ids=_requested_viewer_role_ids(request.form),
            user_ids=_requested_viewer_user_ids(request.form),
        )
        _store_manager_scope(
            new_task,
            request.form.get("manager_scope_mode", "none"),
            role_ids=_requested_manager_role_ids(request.form),
            user_ids=_requested_manager_user_ids(request.form),
        )
        db.session.add(new_task)
        db.session.flush()

        if task_mode in {"FILE", "FORM"}:
            _create_assignment_records(
                new_task,
                assignees,
                assign_type=request.form.get("assign_type", "unit"),
                title_snapshot=new_task.title,
            )

        if task_mode == "OUTLINE":
            outline_item_configs = _parse_outline_item_configs_from_request(request.form)
            if not outline_item_configs and not blueprint_items:
                bulk_titles = _parse_bulk_child_task_titles(
                    request.form.get("bulk_titles") or request.form.get("bulk_items")
                )
                outline_file = request.files.get("outline_file")
                if outline_file and outline_file.filename and not bulk_titles:
                    try:
                        bulk_titles.extend(_parse_outline_upload_titles(outline_file))
                    except ValueError:
                        bulk_titles = []
                if bulk_titles:
                    child_report_kind = str(request.form.get("child_report_kind") or "narrative").strip().lower()
                    if child_report_kind not in CHILD_TASK_ALLOWED_REPORT_KINDS:
                        child_report_kind = "narrative"
                    attachment_required = _report_checkbox_value(request.form.get("child_attachment_required"))
                    outline_item_configs = [
                        {
                            "title": item_title,
                            "report_kind": child_report_kind,
                            "attachment_required": bool(attachment_required),
                        }
                        for item_title in bulk_titles
                    ]

            if outline_item_configs:
                created_item_by_form_index = {}
                for index, item_config in enumerate(outline_item_configs, start=1):
                    item_content = str(item_config.get("content") or "").strip()
                    number_fields = item_config.get("number_fields") or []
                    guide_text = None
                    if number_fields:
                        try:
                            guide_text = json.dumps(number_fields, ensure_ascii=False)
                        except Exception:
                            guide_text = None
                    parent_item_id = None
                    parent_index = item_config.get("parent_index")
                    if parent_index is not None and parent_index in created_item_by_form_index:
                        parent_item_id = created_item_by_form_index[parent_index]
                    task_item = TaskItem(
                        task_id=new_task.id,
                        parent_item_id=parent_item_id,
                        item_code=str(index),
                        title=item_config["title"],
                        content=item_content or None,
                        guide_text=guide_text,
                        is_required=True,
                        output_type="OUTLINE",
                        report_kind=item_config.get("report_kind") or "narrative",
                        attachment_required=bool(item_config.get("attachment_required")),
                        deadline=new_task.deadline,
                        sort_order=index,
                        report_sources_json=_outline_sources_json(item_config.get("sources") or []),
                    )
                    db.session.add(task_item)
                    db.session.flush()
                    table_cells = item_config.get("table_cells") or {}
                    if table_cells:
                        task_item.table_cells_json = json.dumps(table_cells, ensure_ascii=False)
                        schema = item_config.get("table_schema") or []
                        if schema and not new_task.outline_table_schema_json:
                            new_task.outline_table_schema_json = json.dumps(schema, ensure_ascii=False)
                    if item_config.get("report_secondary") and item_content:
                        linked_item = _find_report_secondary_linked_item(
                            item_content,
                            item_config.get("unit_domains") or [],
                            new_task.id,
                        )
                        if linked_item:
                            task_item.linked_item_id = linked_item.id
                    created_item_by_form_index[item_config["form_index"]] = task_item.id
                    if (
                        item_config.get("inherit")
                        and parent_index is not None
                        and parent_index in created_item_by_form_index
                    ):
                        # Dòng con kế thừa gán từ mục cha: tạo assignment giống cha
                        parent_item = TaskItem.query.filter_by(id=created_item_by_form_index[parent_index]).first()
                        if parent_item:
                            parent_assignments = TaskAssignment.query.filter_by(
                                task_id=new_task.id, task_item_id=parent_item.id
                            ).all()
                            for parent_assignment in parent_assignments:
                                db.session.add(
                                    TaskAssignment(
                                        task_id=new_task.id,
                                        task_item_id=task_item.id,
                                        user_id=parent_assignment.user_id,
                                        assignee_type=parent_assignment.assignee_type,
                                        role_id=parent_assignment.role_id,
                                        title_snapshot=item_config["title"],
                                        status="assigned",
                                        is_required=True,
                                        assigned_at=datetime.now(),
                                    )
                                )
                            continue
                    item_assignees, item_error_message, item_assign_type, item_role_ids = _resolve_outline_item_assignment(
                        item_config, request.form, new_task
                    )
                    if item_error_message:
                        flash(f'Nội dung "{item_config.get("title", "")}": {item_error_message}', "danger")
                        db.session.rollback()
                        return redirect(url_for("tasks_bp.tasks"))
                    _create_assignment_records(
                        new_task,
                        item_assignees,
                        assign_type=item_assign_type,
                        task_item=task_item,
                        title_snapshot=task_item.title,
                        role_id=item_role_ids[0] if len(item_role_ids) == 1 else None,
                    )
            elif blueprint_items:
                for index, item_config in enumerate(blueprint_items, start=1):
                    task_item = TaskItem(
                        task_id=new_task.id,
                        item_code=str(index),
                        title=item_config["title"],
                        content=item_config.get("description"),
                        guide_text=item_config.get("guide_text"),
                        is_required=bool(item_config.get("is_required", True)),
                        output_type="OUTLINE",
                        report_kind=item_config.get("report_kind") or "narrative",
                        attachment_required=bool(item_config.get("attachment_required")),
                        deadline=new_task.deadline,
                        sort_order=item_config.get("sort_order", index - 1),
                    )
                    db.session.add(task_item)
                    db.session.flush()
                    _create_assignment_records(
                        new_task,
                        assignees,
                        assign_type=request.form.get("assign_type", "unit"),
                        task_item=task_item,
                        title_snapshot=task_item.title,
                    )

        if task_mode == "FORM":
            field_defs = google_form_field_defs if form_provider == "google" else _parse_task_form_fields_from_request(request.form)
            if not field_defs and workflow_blueprint:
                field_defs = workflow_blueprint_form_field_defs(workflow_blueprint)
            if not field_defs and form_provider != "google":
                flash("Cần cấu hình ít nhất một trường dữ liệu cho biểu mẫu.", "danger")
                db.session.rollback()
                return redirect(url_for("tasks_bp.tasks"))
            try:
                _validate_task_visibility_before_publish(
                    "FORM",
                    assignees,
                    assign_type=assign_type,
                    domain=domain,
                    role_ids=assign_role_ids,
                    user_ids=assign_user_ids,
                    field_defs=field_defs,
                    ignored_form_field_labels=[google_form_match_field] if form_provider == "google" and google_form_match_field else [],
                )
            except ValueError as visibility_error:
                flash(str(visibility_error), "danger")
                db.session.rollback()
                return redirect(url_for("tasks_bp.tasks"))
            for field_def in field_defs:
                db.session.add(TaskFormField(task_id=new_task.id, **_task_form_field_db_kwargs(field_def)))
        elif task_mode == "FILE":
            try:
                _validate_task_visibility_before_publish(
                    "FILE",
                    assignees,
                    assign_type=assign_type,
                    domain=domain,
                    role_ids=assign_role_ids,
                    user_ids=assign_user_ids,
                    report_schema=report_schema,
                )
            except ValueError as visibility_error:
                flash(str(visibility_error), "danger")
                db.session.rollback()
                return redirect(url_for("tasks_bp.tasks"))

        db.session.commit()

        for user in assignees:
            push_notif(user.id, "Công việc mới", f"Bạn vừa được giao: {new_task.title}", f"/tasks/{new_task.id}")

        # Send email notifications to assigned users
        try:
            protocol = request.scheme
            host = request.host
            base_url = f"{protocol}://{host}"
            email_result = send_task_assignment_emails(assignees, new_task, base_url=base_url)
            if email_result.get("skipped"):
                for uid, reason in email_result["skipped"]:
                    logger.warning(f"Email skipped for user {uid}: {reason}")
        except Exception as e:
            logger.error(f"Failed to send task assignment emails: {e}")

        flash("Đã tạo công việc mới.", "success")
        return redirect(url_for("tasks_bp.task_detail", tid=new_task.id))

    candidate_tasks = (
        Task.query.options(joinedload(Task.assignments).joinedload(TaskAssignment.user))
        .filter(Task.parent_task_id.is_(None))
        .order_by(Task.created_at.desc(), Task.id.desc())
        .all()
    )

    visible_tasks = []
    for task in candidate_tasks:
        is_executor = TaskAssignment.query.filter_by(task_id=task.id, user_id=session["uid"]).first() is not None
        is_manager = _can_manage_task(task, user=current_user)
        is_viewer = _can_watch_task(task, user=current_user)
        if not task_visible_for_user(
            task,
            session["uid"],
            can_view_all_tasks=can_view_all_tasks,
            is_admin=is_admin,
            is_executor=is_executor,
            is_manager=is_manager,
            is_viewer=is_viewer,
        ):
            continue

        sync_record_categories([task], task_fields, attr_name="category", prefer_stable=True)
        sync_record_categories([task], pro_units, attr_name="domain", prefer_stable=True)
        sync_record_categories([task], task_types, attr_name="task_type", prefer_stable=True)
        sync_record_categories([task], priority_items, attr_name="priority", prefer_stable=True)
        _decorate_task_categories(task, task_fields, pro_units, task_types, priority_items)
        visible_tasks.append(
            prepare_task_workspace_record(
                task,
                session["uid"],
                is_lead,
                _build_rebuilt_task_summary,
                _task_mode,
                _task_mode_label,
                _task_mode_description,
                _task_assignment_status_label,
                _can_edit_task,
                _can_delete_task,
                _task_assignment_display_status,
                build_task_workspace_attrs,
                today=datetime.now().date(),
            )
        )

    list_context = build_task_list_page_context(visible_tasks, TASK_MODE_DEFAULT)
    current_task_view = (request.args.get("view") or "attention").strip().lower()
    sidebar_submenu_items = []
    if list_context["attention_tasks"]:
        sidebar_submenu_items.append({
            "label": "Cần xử lý ngay",
            "href": url_for("tasks_bp.tasks", view="attention") + "#attention-tasks",
            "count": len(list_context["attention_tasks"]),
            "active": current_task_view == "attention",
        })
    sidebar_submenu_items.extend([
        {
            "label": "Việc của tôi",
            "href": url_for("tasks_bp.tasks", view="my") + "#my-tasks",
            "count": len(list_context["my_tasks"]),
            "active": current_task_view == "my",
        },
        {
            "label": "Tôi giao / theo dõi",
            "href": url_for("tasks_bp.tasks", view="managed") + "#managed-tasks",
            "count": len(list_context["managed_tasks"]),
            "active": current_task_view == "managed",
        },
        {
            "label": "Chỉ xem / tra cứu",
            "href": url_for("tasks_bp.tasks", view="watch") + "#watch-tasks",
            "count": len(list_context["watch_tasks"]),
            "active": current_task_view == "watch",
        },
    ])
    if is_lead or is_admin:
        sidebar_submenu_items.append({
            "label": "Báo cáo định kỳ",
            "href": url_for("tasks_bp.report_dashboard"),
            "count": None,
            "active": False,
        })

    return render_template(
        "tasks_rebuild.html",
        tasks=list_context["tasks"],
        attention_tasks=list_context["attention_tasks"],
        my_tasks=list_context["my_tasks"],
        managed_tasks=list_context["managed_tasks"],
        watch_tasks=list_context["watch_tasks"],
        outline_tasks=list_context["outline_tasks"],
        file_tasks=list_context["file_tasks"],
        form_tasks=list_context["form_tasks"],
        users=active_users,
        roles=roles,
        pro_units=stable_form_category_options(pro_units),
        task_fields=task_fields,
        task_types=stable_form_category_options(task_types),
        priority_items=stable_form_category_options(priority_items),
        is_lead=is_lead,
        is_admin=is_admin,
        stats=list_context["stats"],
        workflow_blueprint_examples=workflow_blueprint_example_catalog(),
        sidebar_submenu_parent="tasks",
        sidebar_submenu_title="Công việc",
        sidebar_submenu_items=sidebar_submenu_items,
    )

def _task_detail_v2(tid):
    task = Task.query.options(joinedload(Task.assignments).joinedload(TaskAssignment.user)).filter_by(id=tid).first()
    if not task:
        return "Not Found", 404

    perms = _current_perms()
    can_view_all_tasks = _can_view_all_tasks(perms)
    is_lead = _can_process_task_module(perms)
    is_admin = bool(current_is_admin())
    current_user = db.session.get(User, session["uid"])
    is_executor = TaskAssignment.query.filter_by(task_id=task.id, user_id=session["uid"]).first() is not None
    can_manage_task_view = bool(is_admin or is_lead or _can_edit_task(task) or _can_manage_task(task, user=current_user))
    can_watch_task_view = bool(_can_watch_task(task, user=current_user))

    if not (can_view_all_tasks or can_manage_task_view or can_watch_task_view or is_executor or task.author_id == session["uid"]):
        flash("Bạn không có quyền xem công việc này.", "danger")
        return redirect(url_for("tasks_bp.tasks"))

    _lazy_repair_task_runtime(task, include_children=False, commit=True)

    pro_units = _task_domain_options()
    task_fields = _task_field_options()
    task_types = _task_type_options()
    priority_items = _task_priority_options()
    sync_record_categories([task], task_fields, attr_name="category", prefer_stable=True)
    sync_record_categories([task], pro_units, attr_name="domain", prefer_stable=True)
    sync_record_categories([task], task_types, attr_name="task_type", prefer_stable=True)
    sync_record_categories([task], priority_items, attr_name="priority", prefer_stable=True)
    _decorate_task_categories(task, task_fields, pro_units, task_types, priority_items)

    mode = _task_mode(task)
    setattr(task, "task_mode", mode)
    setattr(task, "task_mode_label", _task_mode_label(mode))
    setattr(task, "task_mode_description", _task_mode_description(mode))
    if mode == "FORM" and str(getattr(task, "form_provider", "") or "").strip().lower() == "google":
        _apply_task_google_form_view_state(task)

    detail_page_context = build_task_detail_page_context(
        task,
        session["uid"],
        mode,
        can_manage_task_view,
        is_executor,
        _build_rebuilt_task_summary,
        _parse_outline_item_rows,
        _build_outline_group_rows,
        _build_file_task_rows,
        _build_form_task_rows,
        _task_form_field_views,
        _task_detail_context,
    )
    if mode == "OUTLINE" and not (can_manage_task_view or can_watch_task_view):
        detail_page_context["outline_groups"] = _filter_outline_groups_for_executor_scope(detail_page_context["outline_groups"])
        visible_item_ids = {
            getattr(row.get("item"), "id", None)
            for group in detail_page_context["outline_groups"]
            for row in (group.get("rows") or [])
        }
        detail_page_context["outline_rows"] = [
            row for row in (detail_page_context["outline_rows"] or [])
            if getattr(row.get("item"), "id", None) in visible_item_ids
        ]
    elif mode == "FILE" and not (can_manage_task_view or can_watch_task_view):
        detail_page_context["file_rows"] = _filter_assignment_rows_for_executor_scope(
            detail_page_context["file_rows"],
            detail_page_context["my_file_assignment"],
        )
    elif mode == "FORM" and not (can_manage_task_view or can_watch_task_view):
        detail_page_context["form_rows"] = _filter_assignment_rows_for_executor_scope(
            detail_page_context["form_rows"],
            detail_page_context["my_form_assignment"],
        )
    file_report_comments = TaskComment.query.filter_by(task_id=task.id).order_by(TaskComment.created_at.asc(), TaskComment.id.asc()).all() if mode == "FILE" else []
    my_file_report_form = _build_structured_task_report_form(task, detail_page_context["my_file_assignment"], current_user) if mode == "FILE" and detail_page_context["my_file_assignment"] else None
    my_form_field_views = _task_form_field_views_for_user(task, current_user) if mode == "FORM" and detail_page_context["my_form_assignment"] else []
    if mode == "FILE":
        for row in detail_page_context["file_rows"]:
            row["report_context"] = _build_assignment_report_context(row["assignment"], file_report_comments, task=task)
    file_progress_groups = _task_assignment_progress_groups(detail_page_context["file_rows"]) if mode == "FILE" else {"unit_cards": [], "role_groups": []}
    form_progress_groups = _task_assignment_progress_groups(detail_page_context["form_rows"]) if mode == "FORM" else {"unit_cards": [], "role_groups": []}
    delivery_contract_rows = []
    if mode == "FILE":
        delivery_contract_rows = detail_page_context["file_rows"]
        if is_executor and detail_page_context["my_file_assignment"]:
            delivery_contract_rows = _filter_assignment_rows_for_executor_scope(
                delivery_contract_rows,
                detail_page_context["my_file_assignment"],
            )
    elif mode == "FORM":
        delivery_contract_rows = detail_page_context["form_rows"]
        if is_executor and detail_page_context["my_form_assignment"]:
            delivery_contract_rows = _filter_assignment_rows_for_executor_scope(
                delivery_contract_rows,
                detail_page_context["my_form_assignment"],
            )
    delivery_contract_groups = _task_delivery_contract_groups(task, mode, delivery_contract_rows) if mode in {"FILE", "FORM"} else []
    active_users = []
    roles = []
    if can_manage_task_view:
        active_users = User.query.filter_by(is_active=True).order_by(User.unit_area.asc(), User.fullname.asc()).all()
        active_users = apply_reference_display(
            sync_record_categories(
                active_users,
                module_category_options("contacts", "unit_name", "Đơn vị"),
                attr_name="unit_area",
                prefer_stable=True,
            ),
            "unit_area",
            module_category_options("contacts", "unit_name", "Đơn vị"),
            display_attr="unit_area_display",
            fallback_label="Chưa có đơn vị",
        )
        roles = AppRole.query.order_by(AppRole.name.asc()).all()
    outline_import_preview_rows = _get_outline_import_preview(task.id) if mode == "OUTLINE" and can_manage_task_view else []
    outline_matrix = _build_outline_progress_matrix(task, session["uid"]) if mode == "OUTLINE" else None

    return render_template(
        "task_detail_rebuild.html",
        task=task,
        pro_units=stable_form_category_options(pro_units),
        task_fields=stable_form_category_options(task_fields),
        task_types=stable_form_category_options(task_types),
        priority_items=stable_form_category_options(priority_items),
        can_edit_task=_can_edit_task(task),
        can_delete_task=_can_delete_task(task, is_lead=is_lead),
        can_manage_task_view=can_manage_task_view,
        can_watch_task_view=can_watch_task_view,
        can_submit=is_executor,
        is_lead=is_lead,
        is_admin=is_admin,
        users=active_users,
        roles=roles,
        outline_rows=detail_page_context["outline_rows"],
        outline_groups=detail_page_context["outline_groups"],
        outline_import_preview_rows=outline_import_preview_rows,
        outline_matrix=outline_matrix,
        file_rows=detail_page_context["file_rows"],
        file_assignment_unit_cards=file_progress_groups["unit_cards"],
        file_assignment_role_groups=file_progress_groups["role_groups"],
        delivery_contract_groups=delivery_contract_groups,
        my_file_report_form=my_file_report_form,
        form_fields=detail_page_context["form_fields"],
        form_field_views=detail_page_context["form_field_views"],
        form_rows=detail_page_context["form_rows"],
        form_assignment_unit_cards=form_progress_groups["unit_cards"],
        form_assignment_role_groups=form_progress_groups["role_groups"],
        my_file_assignment=detail_page_context["my_file_assignment"],
        my_file_submission=detail_page_context["my_file_submission"],
        my_form_assignment=detail_page_context["my_form_assignment"],
        my_form_submission=detail_page_context["my_form_submission"],
        my_form_payload=detail_page_context["my_form_payload"],
        my_form_field_views=my_form_field_views,
        summary=detail_page_context["summary"],
        detail_context=detail_page_context["detail_context"],
        status_labels=TASK_ASSIGNMENT_STATUS_LABELS,
        report_period=_task_report_period(task),
        report_kind_label=_task_report_kind_label(task),
        current_cycle=_task_current_cycle(task),
    )

def _create_outline_items_v2(tid):
    parent_task = Task.query.filter_by(id=tid).first()
    if not parent_task:
        return "Not Found", 404

    if _task_mode(parent_task) != "OUTLINE":
        flash("Công việc này không dùng chế độ đề cương.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    if not _can_edit_task(parent_task):
        flash("Bạn không có quyền thêm đầu mục cho công việc này.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    item_configs = _parse_outline_item_configs_from_request(request.form)
    bulk_titles = _parse_bulk_child_task_titles(request.form.get("bulk_titles") or request.form.get("bulk_items"))
    outline_file = request.files.get("outline_file")
    if outline_file and outline_file.filename and not item_configs:
        try:
            bulk_titles.extend(_parse_outline_upload_titles(outline_file))
        except ValueError as outline_error:
            flash(str(outline_error), "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        bulk_titles = _parse_bulk_child_task_titles("\n".join(bulk_titles))

    if not item_configs:
        child_report_kind = str(request.form.get("child_report_kind") or "narrative").strip().lower()
        if child_report_kind not in CHILD_TASK_ALLOWED_REPORT_KINDS:
            child_report_kind = "narrative"
        attachment_required = _report_checkbox_value(request.form.get("child_attachment_required"))
        item_configs = [
            {
                "title": item_title,
                "report_kind": child_report_kind,
                "attachment_required": bool(attachment_required),
            }
            for item_title in bulk_titles
        ]

    if not item_configs:
        flash("Cần tạo ít nhất một nội dung báo cáo trước khi gán.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    current_count = TaskItem.query.filter_by(task_id=parent_task.id).count()

    created_item_by_form_index = {}
    for index, item_config in enumerate(item_configs, start=1):
        item_content = str(item_config.get("content") or "").strip()
        number_fields = item_config.get("number_fields") or []
        guide_text = None
        if number_fields:
            try:
                guide_text = json.dumps(number_fields, ensure_ascii=False)
            except Exception:
                guide_text = None
        parent_item_id = None
        parent_index = item_config.get("parent_index")
        if parent_index is not None and parent_index in created_item_by_form_index:
            parent_item_id = created_item_by_form_index[parent_index]
        task_item = TaskItem(
            task_id=parent_task.id,
            parent_item_id=parent_item_id,
            item_code=str(current_count + index),
            title=item_config["title"],
            content=item_content or None,
            guide_text=guide_text,
            is_required=True,
            output_type="OUTLINE",
            report_kind=item_config["report_kind"],
            attachment_required=bool(item_config["attachment_required"]),
            deadline=parent_task.deadline,
            sort_order=current_count + index,
        )
        db.session.add(task_item)
        db.session.flush()
        created_item_by_form_index[item_config.get("form_index", index - 1)] = task_item.id
        if (
            item_config.get("inherit")
            and parent_index is not None
            and parent_index in created_item_by_form_index
        ):
            # Dòng con kế thừa gán từ mục cha: tạo assignment giống cha
            parent_item = TaskItem.query.filter_by(id=created_item_by_form_index[parent_index]).first()
            if parent_item:
                parent_assignments = TaskAssignment.query.filter_by(
                    task_id=parent_task.id, task_item_id=parent_item.id
                ).all()
                for parent_assignment in parent_assignments:
                    db.session.add(
                        TaskAssignment(
                            task_id=parent_task.id,
                            task_item_id=task_item.id,
                            user_id=parent_assignment.user_id,
                            assignee_type=parent_assignment.assignee_type,
                            role_id=parent_assignment.role_id,
                            title_snapshot=item_config["title"],
                            status="assigned",
                            is_required=True,
                            assigned_at=datetime.now(),
                        )
                    )
                continue
        assignees, error_message, assign_type, role_ids = _resolve_outline_item_assignment(item_config, request.form, parent_task)
        if error_message:
            flash(f'Nội dung "{item_config["title"]}": {error_message}', "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        _create_assignment_records(
            parent_task,
            assignees,
            assign_type=assign_type,
            task_item=task_item,
            title_snapshot=item_config["title"],
            role_id=role_ids[0] if len(role_ids) == 1 else None,
        )

    db.session.commit()
    _clear_outline_import_preview(parent_task.id)
    flash(f"Đã thêm {len(item_configs)} đầu mục.", "success")
    return redirect(url_for("tasks_bp.task_detail", tid=tid))

def _preview_outline_import_v2(tid):
    parent_task = Task.query.filter_by(id=tid).first()
    if not parent_task:
        return "Not Found", 404

    if _task_mode(parent_task) != "OUTLINE":
        flash("Công việc này không dùng chế độ đề cương.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    if not _can_edit_task(parent_task):
        flash("Bạn không có quyền nạp đề cương cho công việc này.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    outline_file = request.files.get("outline_file")
    if not outline_file or not outline_file.filename:
        flash("Cần chọn file đề cương trước khi nạp.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    try:
        parsed_rows = _parse_outline_upload_rows(outline_file)
    except ValueError as outline_error:
        flash(str(outline_error), "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    if not parsed_rows:
        flash("Không tìm thấy đầu mục hợp lệ trong file đề cương.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    default_report_kind = str(request.form.get("child_report_kind") or "narrative").strip().lower()
    if default_report_kind not in CHILD_TASK_ALLOWED_REPORT_KINDS:
        default_report_kind = "narrative"
    attachment_required = _report_checkbox_value(request.form.get("child_attachment_required"))

    preview_rows = []
    detected_count = 0
    for row in parsed_rows:
        assignee_detected = bool(row.get("assignee_detected"))
        preview_rows.append(
            {
                "title": row["title"],
                "content": row.get("content") or "",
                "heading": row.get("heading") or "",
                "parent_row_index": row.get("parent_row_index"),
                "report_kind": default_report_kind,
                "attachment_required": bool(attachment_required),
                "assign_type": row.get("assign_type") or "",
                "domain": row.get("domain") or "",
                "unit_domains": row.get("unit_domains") or [],
                "role_ids": row.get("role_ids") or [],
                "user_ids": row.get("user_ids") or [],
                "assignee_hint": row.get("assignee_hint") or "",
            }
        )
        if assignee_detected:
            detected_count += 1
    _set_outline_import_preview(parent_task.id, preview_rows)
    if detected_count:
        flash(
            f"Đã nạp {len(preview_rows)} nội dung từ đề cương; tự nhận diện người nhận cho {detected_count} đầu mục. "
            "Kiểm tra cột 'Người nhận' rồi bấm Tạo khi chính xác.",
            "success",
        )
    else:
        flash(
            f"Đã nạp {len(preview_rows)} nội dung từ đề cương. Chưa phát hiện 'giao cho ai' trong file, "
            "bạn gán người nhận ở bước 2 trước khi tạo.",
            "info",
        )
    return redirect(url_for("tasks_bp.task_detail", tid=tid))

def _update_task_status_v2(tid):
    task = Task.query.filter_by(id=tid).first()
    if not task:
        return "Not Found", 404

    task_item_id = request.form.get("task_item_id", "").strip()
    query = TaskAssignment.query.filter_by(task_id=tid, user_id=session["uid"])
    if task_item_id.isdigit():
        query = query.filter_by(task_item_id=int(task_item_id))
    else:
        query = query.filter(TaskAssignment.task_item_id.is_(None))
    assignment = query.first()
    if not assignment:
        flash("Bạn không được giao nội dung này.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    assignment.status = "in_progress"
    assignment.updated_at = datetime.now()
    db.session.commit()
    flash("Đã tiếp nhận công việc.", "success")
    return redirect(url_for("tasks_bp.task_detail", tid=tid))

def _submit_task_report_v2(tid):
    task = Task.query.filter_by(id=tid).first()
    if not task:
        return "Not Found", 404

    mode = _task_mode(task)
    item = None
    query = TaskAssignment.query.filter_by(task_id=tid, user_id=session["uid"])
    task_item_id = request.form.get("task_item_id", "").strip()
    if mode == "OUTLINE":
        if not task_item_id.isdigit():
            flash("Thiếu đầu mục cần báo cáo.", "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        item = TaskItem.query.filter_by(id=int(task_item_id), task_id=tid).first()
        if not item:
            flash("Không tìm thấy đầu mục cần báo cáo.", "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        query = query.filter_by(task_item_id=item.id)
    else:
        query = query.filter(TaskAssignment.task_item_id.is_(None))

    assignment = query.first()
    if not assignment:
        flash("Bạn không được giao nội dung này.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    narrative = (request.form.get("report_content") or request.form.get("report_narrative") or "").strip()
    report_file = request.files.get("report_file")
    numeric_value = None
    payload = {}
    structured_submission = None
    current_user = db.session.get(User, session["uid"])

    if mode == "OUTLINE" and item and item.report_kind == "number":
        # Nhận số liệu theo từng ô trống (report_number_value_<blank_id>)
        per_field_values = {}
        for field_key, raw_field in request.form.items():
            if field_key.startswith("report_number_value_"):
                field_idx = field_key[len("report_number_value_"):]
                field_text = str(raw_field or "").strip()
                if field_text:
                    parsed = _parse_outline_blank_value(field_text)
                    if parsed is None:
                        flash("Số liệu không hợp lệ.", "danger")
                        return redirect(url_for("tasks_bp.task_detail", tid=tid))
                    per_field_values[field_idx] = parsed
        raw_value = (request.form.get("report_number") or "").strip()
        if per_field_values:
            if not raw_value:
                raw_value = str(next(iter(per_field_values.values())))
            payload["values"] = per_field_values
            numeric_value = _outline_blank_numeric(raw_value)
            payload["reported_value"] = numeric_value
        else:
            if not raw_value:
                flash("Cần nhập số liệu cho đầu mục này.", "danger")
                return redirect(url_for("tasks_bp.task_detail", tid=tid))
            numeric_value = _outline_blank_numeric(raw_value)
            if numeric_value is None:
                flash("Số liệu không hợp lệ.", "danger")
                return redirect(url_for("tasks_bp.task_detail", tid=tid))
            payload["reported_value"] = numeric_value

    if mode == "FORM":
        missing_labels = []
        visible_form_fields = _task_form_fields_for_user(task, current_user)
        if not visible_form_fields:
            flash("Bạn chưa được giao trường dữ liệu nào trong biểu mẫu này.", "warning")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        for field in visible_form_fields:
            field_key = f"form_field_{field.field_key}"
            field_type = _normalize_task_form_field_type(field.field_type)
            if field_type == "checkbox":
                value = request.form.getlist(field_key)
            else:
                value = (request.form.get(field_key) or "").strip()
            if field_type == "number" and not _task_form_value_is_empty(value):
                try:
                    if isinstance(value, str):
                        value = float(value.replace(",", ""))
                except ValueError:
                    flash(f"Trường {field.field_label} phải là số hợp lệ.", "danger")
                    return redirect(url_for("tasks_bp.task_detail", tid=tid))
            if field_type == "table" and isinstance(value, str):
                table_rows = []
                for line in value.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    table_rows.append([cell.strip() for cell in line.split("|")])
                value = table_rows
            if getattr(field, "is_required", False) and _task_form_value_is_empty(value):
                missing_labels.append(field.field_label)
            payload[field.field_key] = value
        if missing_labels:
            flash("Cần điền các trường bắt buộc: " + ", ".join(missing_labels) + ".", "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        narrative = ""
    elif mode == "FILE":
        try:
            structured_submission = _parse_structured_file_report_submission(
                task,
                assignment,
                current_user,
                request.form,
                report_file,
            )
        except ValueError as exc:
            flash(str(exc), "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))
        if structured_submission:
            narrative = structured_submission["narrative"]
            numeric_value = structured_submission["numeric_value"]
            payload = structured_submission["payload"]

    if mode != "FORM" and not narrative and not report_file and numeric_value is None:
        if not structured_submission:
            flash("Cần nhập nội dung hoặc đính kèm tệp báo cáo.", "danger")
            return redirect(url_for("tasks_bp.task_detail", tid=tid))

    submission = TaskSubmission(
        task_id=task.id,
        task_item_id=getattr(item, "id", None),
        assignment_id=assignment.id,
        submitted_by=session["uid"],
        submission_type=mode,
        status="submitted",
        narrative_content=narrative or None,
        numeric_value=numeric_value,
        payload_json=json.dumps(payload, ensure_ascii=False) if payload else None,
        submitted_at=datetime.now(),
    )
    current_cycle = _task_current_cycle(task)
    if current_cycle:
        submission.cycle_key = str(current_cycle.get("key") or "")[:50] or None
        submission.cycle_label = str(current_cycle.get("label") or "")[:100] or None
    db.session.add(submission)
    db.session.flush()

    if report_file and report_file.filename:
        file_meta = _store_uploaded_task_file(report_file, task.id, assignment.id, prefix="submission")
        if file_meta:
            db.session.add(
                TaskSubmissionFile(
                    submission_id=submission.id,
                    original_name=file_meta["original_name"],
                    stored_name=file_meta["stored_name"],
                    stored_path=file_meta["stored_path"],
                    file_ext=file_meta["file_ext"],
                    mime_type=file_meta["mime_type"],
                    file_size=file_meta["file_size"],
                )
            )
            submission.attachment_name = file_meta["original_name"]
            submission.attachment_path = file_meta["stored_path"]
            assignment.result_file = file_meta["stored_name"]
            if structured_submission:
                payload["attachment_name"] = file_meta["original_name"]

    assignment.status = "submitted"
    assignment.submitted_at = datetime.now()
    assignment.last_submission_id = submission.id
    assignment.report_payload_json = json.dumps(
        {
            "mode": payload.get("mode") if isinstance(payload, dict) else None,
            "narrative": narrative,
            "numeric_value": numeric_value,
            "payload": payload,
            "values": payload.get("values", {}) if isinstance(payload, dict) else {},
            "attachment_name": payload.get("attachment_name", "") if isinstance(payload, dict) else "",
            "submitted_at": submission.submitted_at.strftime("%d/%m/%Y %H:%M"),
        },
        ensure_ascii=False,
    )
    assignment.updated_at = datetime.now()
    _sync_assignment_group_submission(
        task,
        assignment,
        submission,
        report_payload_json=assignment.report_payload_json or "",
        result_file=assignment.result_file or "",
        submitted_at=assignment.submitted_at,
        updated_at=assignment.updated_at,
        status="submitted",
    )

    comment_text = narrative or ('Đã cập nhật biểu mẫu báo cáo' if structured_submission else ('Đã nộp biểu mẫu' if mode == 'FORM' else 'Đã nộp báo cáo'))
    if structured_submission:
        comment_text = _build_structured_task_report_comment(_load_task_report_schema(task), payload)
    db.session.add(
        TaskComment(
            task_id=task.id,
            user_id=session["uid"],
            user_name=session.get("fullname", "Người dùng"),
            content=f"[BÁO CÁO] {comment_text}",
        )
    )
    _propagate_submission_to_linked_items(task, item, assignment, submission)
    db.session.commit()
    flash("Đã gửi báo cáo.", "success")
    return redirect(url_for("tasks_bp.task_detail", tid=tid))

def _export_form_task_v2(tid):
    if Workbook is None:
        flash("Máy chủ chưa cài thư viện xuất Excel.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    task = Task.query.filter_by(id=tid).first()
    if not task:
        return "Not Found", 404

    current_user = db.session.get(User, session["uid"])
    perms = _current_perms()
    can_manage_task_view = bool(
        bool(current_is_admin())
        or _can_process_task_module(perms)
        or _can_edit_task(task)
        or _can_manage_task(task, user=current_user)
        or _can_watch_task(task, user=current_user)
    )
    if not can_manage_task_view:
        flash("Bạn không có quyền xuất dữ liệu biểu mẫu.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    fields, rows = _build_form_task_rows(task, session["uid"])
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Du lieu"
    headers = ["Đơn vị", "Người nhận", "Trạng thái", "Thời điểm nộp"]
    headers.extend([field.field_label for field in fields])
    sheet.append(headers)

    for row in rows:
        user = getattr(row["assignment"], "user", None)
        user_name = getattr(user, "fullname", None) or getattr(user, "username", None) or ""
        unit_name = _task_assignee_unit_name(user) if user else ""
        submission = row["submission"]
        payload = row["payload"] or {}
        data_row = [
            unit_name,
            user_name,
            _task_assignment_status_label(row["assignment"].status),
            submission.submitted_at.strftime("%d/%m/%Y %H:%M") if submission and submission.submitted_at else "",
        ]
        for field in fields:
            value = payload.get(field.field_key, "")
            if isinstance(value, list):
                if value and isinstance(value[0], list):
                    value = " || ".join(" | ".join(str(cell) for cell in item) for item in value)
                else:
                    value = ", ".join(str(item) for item in value)
            data_row.append(value)
        sheet.append(data_row)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    safe_name = secure_filename(task.title or f"task_{task.id}") or f"task_{task.id}"
    return send_file(
        output,
        as_attachment=True,
        download_name=f"du_lieu_bieu_mau_{safe_name}_{task.id}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

def _build_outline_progress_matrix(task, current_uid):
    """Ma trận tiến độ: hàng = đầu mục, cột = đơn vị nhận việc."""
    rows = _parse_outline_item_rows(task, current_uid)
    unit_names = []
    for row in rows:
        for assignment in row["assignments"]:
            user = getattr(assignment, "user", None) or db.session.get(User, getattr(assignment, "user_id", None))
            unit_name = _task_assignee_unit_name(user)
            if unit_name not in unit_names:
                unit_names.append(unit_name)
    unit_names.sort(key=lambda name: remove_accents(name).lower())

    matrix_rows = []
    for row in rows:
        item = row["item"]
        cells = []
        item_submitted = 0
        item_total = len(row["assignments"])
        for unit_name in unit_names:
            unit_assignments = []
            for assignment in row["assignments"]:
                user = getattr(assignment, "user", None) or db.session.get(User, getattr(assignment, "user_id", None))
                if _task_assignee_unit_name(user) == unit_name:
                    unit_assignments.append(assignment)
            unit_submitted = sum(1 for assignment in unit_assignments if _task_is_submitted(assignment))
            item_submitted += unit_submitted
            cell_numbers = []
            for assignment in unit_assignments:
                submission = row["latest_submissions"].get(assignment.id)
                if not submission or item.report_kind != "number":
                    continue
                values = _outline_submission_values(submission)
                first_value = next(iter(values.values()), None)
                cell_numbers.append(
                    {
                        "unit_name": unit_name,
                        "values": values,
                        "first_value": first_value,
                        "numeric": _outline_blank_numeric(first_value),
                        "submitted": _task_is_submitted(assignment),
                    }
                )
            cells.append(
                {
                    "unit_name": unit_name,
                    "submitted_count": unit_submitted,
                    "total_count": len(unit_assignments),
                    "done": bool(unit_assignments) and unit_submitted >= len(unit_assignments),
                    "numbers": cell_numbers,
                    "assignments": [
                        {
                            "assignment": assignment,
                            "status": assignment.status,
                            "status_label": _task_assignment_status_label(assignment.status),
                            "status_class": _task_assignment_status_class(assignment.status),
                            "submitted": _task_is_submitted(assignment),
                            "submission": row["latest_submissions"].get(assignment.id),
                        }
                        for assignment in unit_assignments
                    ],
                }
            )
        aggregate_total = None
        aggregate_count = 0
        if item.report_kind == "number":
            numeric_values = []
            for cell in cells:
                for number in cell["numbers"]:
                    if number.get("numeric") is not None:
                        numeric_values.append(number["numeric"])
            aggregate_count = len(numeric_values)
            if numeric_values:
                aggregate_total = sum(numeric_values)
        matrix_rows.append(
            {
                "item": item,
                "cells": cells,
                "submitted_count": item_submitted,
                "total_count": item_total,
                "done": item_total > 0 and item_submitted >= item_total,
                "percent": round(item_submitted / item_total * 100) if item_total else 0,
                "aggregate_total": aggregate_total,
                "aggregate_count": aggregate_count,
            }
        )

    total_submitted = sum(matrix_row["submitted_count"] for matrix_row in matrix_rows)
    total_count = sum(matrix_row["total_count"] for matrix_row in matrix_rows)
    return {
        "unit_names": unit_names,
        "rows": matrix_rows,
        "total_submitted": total_submitted,
        "total_count": total_count,
        "percent": round(total_submitted / total_count * 100) if total_count else 0,
    }

def _export_outline_word_v2(tid):
    task = Task.query.filter_by(id=tid).first()
    if not task:
        return "Not Found", 404

    if _task_mode(task) != "OUTLINE":
        flash("Công việc này không phải dạng báo cáo văn bản theo đề cương.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    current_user = db.session.get(User, session["uid"])
    perms = _current_perms()
    can_manage_task_view = bool(
        bool(current_is_admin())
        or _can_process_task_module(perms)
        or _can_edit_task(task)
        or _can_manage_task(task, user=current_user)
        or _can_watch_task(task, user=current_user)
    )
    if not can_manage_task_view:
        flash("Bạn không có quyền xuất báo cáo tổng hợp.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    if DocxDocument is None:
        flash("Máy chủ chưa cài thư viện tạo file Word.", "danger")
        return redirect(url_for("tasks_bp.task_detail", tid=tid))

    rows = _parse_outline_item_rows(task, session["uid"])
    table_schema_map = _outline_table_schema_map(task)
    document = DocxDocument()
    document.add_heading(str(task.title or f"Công việc #{task.id}"), level=0)

    meta_parts = []
    if task.author_name:
        meta_parts.append(f"Đơn vị giao việc: {task.author_name}")
    if task.deadline:
        meta_parts.append(f"Hạn nộp: {task.deadline.strftime('%d/%m/%Y')}")
    if task.priority:
        meta_parts.append(f"Ưu tiên: {task.priority}")
    if meta_parts:
        meta_paragraph = document.add_paragraph()
        meta_run = meta_paragraph.add_run(" — ".join(meta_parts))
        meta_run.bold = True
    if task.content:
        document.add_paragraph(str(task.content))

    if not rows:
        document.add_paragraph("Chưa có đầu mục nào được thiết lập cho công việc này.")
    for index, row in enumerate(rows, start=1):
        item = row["item"]
        item_code = str(getattr(item, "item_code", None) or index)
        content = str(getattr(item, "content", "") or "")
        document.add_heading(f"{item_code}. {item.title}", level=1)
        # Tái hiện bảng (chỉ các cột được tích hiển thị) nếu đầu mục từ đề cương dạng bảng
        item_table_cells = _outline_item_table_cells(item)
        if table_schema_map and item_table_cells:
            columns = sorted(table_schema_map.values(), key=lambda col: int(col.get("index") or 0))
            columns = [col for col in columns if col.get("visible")]
            if columns:
                outline_table = document.add_table(rows=2, cols=len(columns))
                outline_table.style = "Table Grid"
                for col_index, col in enumerate(columns):
                    outline_table.rows[0].cells[col_index].text = str(col.get("header") or "")
                    value = str(item_table_cells.get(str(col.get("index")), "") or "").strip()
                    if not value and col.get("role") == "content":
                        value = content
                    outline_table.rows[1].cells[col_index].text = value
        if not row["assignments"]:
            document.add_paragraph("Chưa giao đơn vị nào cho đầu mục này.")
        synthesis = _task_item_synthesis_text(item)
        if synthesis:
            document.add_paragraph(synthesis)
        number_fields = _outline_item_number_fields(item)
        submitted_with_values = []
        for assignment in row["assignments"]:
            submission = row["latest_submissions"].get(assignment.id)
            if not submission:
                continue
            user = getattr(assignment, "user", None) or db.session.get(User, getattr(assignment, "user_id", None))
            submitted_with_values.append((assignment, submission, user, _outline_submission_values(submission)))
        # Văn bản tổng hợp: nội dung gốc với số liệu đã nộp ghép vào
        if item.report_kind == "number" and number_fields and submitted_with_values:
            merged_parts = []
            for position, (assignment, submission, user, values) in enumerate(submitted_with_values, start=1):
                unit_name = _task_assignee_unit_name(user)
                merged = _outline_merged_content(content, number_fields, values)
                merged_parts.append(f"Số liệu {position} - {unit_name}: {merged.strip()}")
            merged_paragraph = document.add_paragraph()
            merged_run = merged_paragraph.add_run("\n".join(merged_parts))
            merged_run.bold = True
            # Cộng gộp khi quản trị bật
            if item.allow_aggregate:
                numeric_values = [
                    _outline_blank_numeric(value)
                    for values in (v for _, _, _, v in submitted_with_values)
                    for value in values.values()
                ]
                numeric_values = [value for value in numeric_values if value is not None]
                if numeric_values:
                    aggregate_paragraph = document.add_paragraph()
                    aggregate_run = aggregate_paragraph.add_run(
                        f"Tổng cộng: {sum(numeric_values):,.0f}".replace(",", ".")
                    )
                    aggregate_run.bold = True
        # Đã có văn bản tổng hợp của quản trị -> không lặp lại nội dung từng đơn vị.
        if synthesis:
            continue
        for assignment in row["assignments"]:
            user = getattr(assignment, "user", None) or db.session.get(User, getattr(assignment, "user_id", None))
            unit_name = _task_assignee_unit_name(user)
            user_name = getattr(user, "fullname", None) or getattr(user, "username", None) or "Cán bộ"
            submission = row["latest_submissions"].get(assignment.id)
            status_label = _task_assignment_status_label(assignment.status)
            header = document.add_paragraph()
            header.add_run(f"{unit_name} — {user_name} ({status_label})").bold = True
            if not submission:
                document.add_paragraph("Chưa có nội dung báo cáo.")
                continue
            if item.report_kind == "number" and number_fields:
                values = _outline_submission_values(submission)
                field_lines = []
                for field in number_fields:
                    blank_id = field.get("blank_id")
                    submitted = values.get(str(blank_id), values.get(blank_id, ""))
                    if submitted in (None, ""):
                        submitted = field.get("value", "")
                    field_lines.append(f"- {field.get('label', '')}: {submitted} {field.get('unit', '')}".strip())
                if field_lines:
                    document.add_paragraph("\n".join(field_lines))
            elif item.report_kind == "number" and submission.numeric_value is not None:
                document.add_paragraph(f"Số liệu: {submission.numeric_value:g}")
            if submission.narrative_content:
                document.add_paragraph(str(submission.narrative_content))
            for file in (getattr(submission, "files", None) or []):
                document.add_paragraph(f"File minh chứng: {file.original_name or file.stored_name}")
            if getattr(submission, "submitted_at", None):
                submitted_paragraph = document.add_paragraph()
                submitted_run = submitted_paragraph.add_run(
                    f"(Nộp lúc {submission.submitted_at.strftime('%d/%m/%Y %H:%M')})"
                )
                submitted_run.italic = True

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    safe_name = secure_filename(task.title or f"task_{task.id}") or f"task_{task.id}"
    return send_file(
        output,
        as_attachment=True,
        download_name=f"bao_cao_tong_hop_{safe_name}_{task.id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
