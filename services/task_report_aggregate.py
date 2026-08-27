# -*- coding: utf-8 -*-
"""
Aggregator service for task reports.

Provides cycle-aware aggregation of submissions per TaskItem and per unit.
"""
from collections import defaultdict
from decimal import Decimal
import json
from datetime import datetime

from models import Task, TaskItem, TaskAssignment, TaskSubmission, User
from services.task_modes import _task_assignment_status_label
from services.task_runtime_sync import _task_assignment_rows, _assignment_report_snapshot
from sqlalchemy.orm import joinedload, selectinload


def _get_db():
    """Import lười `app.db` để tránh vòng lặp import app -> routes.tasks ->
    services.task_report_aggregate -> app. Chuỗi này khiến khi chạy
    `python app.py`, app.py bị nạp lần 2 dưới tên module `app` và đăng ký
    blueprint trước khi routes/tasks.py kịp gắn hết decorator (Flask >= 3
    chặn điều này bằng AssertionError)."""
    from app import db

    return db


def _as_number(value):
    try:
        if value is None or value == "":
            return None
        # support numbers stored as strings
        return Decimal(str(value))
    except Exception:
        return None


def build_aggregate_context(task_id, cycle_key=None, group_by="unit"):
    """Build aggregated JSON for a task for the given cycle_key.

    Returns a dict with items, units, and aggregated values per field.
    group_by: 'unit' or 'assignment' (default 'unit')
    """
    db = _get_db()
    task = (
        db.session.query(Task)
        .options(selectinload(Task.task_items))
        .get(task_id)
    )
    if not task:
        return {"error": "task_not_found"}

    # load expected assignments/participants: use TaskAssignment rows for the task scope
    # map units / assignments
    assignment_q = (
        db.session.query(TaskAssignment)
        .filter(TaskAssignment.task_id == task_id)
        .options(selectinload(TaskAssignment.user))
    )
    assignments = assignment_q.all()

    # collect expected unit ids (if unit_id present) or user-based keys
    expected_units = {}
    expected_assignments = {}
    for a in assignments:
        key = None
        if group_by == "unit" and a.unit_id:
            key = f"unit_{a.unit_id}"
            if key not in expected_units:
                expected_units[key] = {"unit_id": a.unit_id, "label": str(a.unit_id)}
        else:
            # fallback to assignment-level grouping (per user)
            key = f"assign_{a.id}"
            expected_assignments[key] = {"assignment_id": a.id, "label": getattr(a.user, "fullname", None) or f"UID{a.user_id}"}

    # load submissions for the task (and optional cycle)
    sub_q = db.session.query(TaskSubmission).filter(TaskSubmission.task_id == task_id)
    if cycle_key:
        sub_q = sub_q.filter(TaskSubmission.cycle_key == cycle_key)
    submissions = sub_q.options(selectinload(TaskSubmission.task_item)).all()

    # index submissions by item_id and by unit/assignment key
    subs_index = defaultdict(lambda: defaultdict(list))
    for s in submissions:
        item_id = s.task_item_id or 0
        # determine key
        key = None
        if group_by == "unit":
            # try assignment -> unit
            if s.assignment and getattr(s.assignment, 'unit_id', None):
                key = f"unit_{s.assignment.unit_id}"
            else:
                key = f"assign_{s.assignment_id or 0}"
        else:
            key = f"assign_{s.assignment_id or 0}"
        subs_index[item_id][key].append(s)

    # prepare items output
    items_out = []
    for item in sorted(task.task_items, key=lambda it: (it.sort_order or 0, it.id)):
        item_id = item.id
        fields = []
        # We support numeric aggregation from payload_json->values or numeric_value
        # determine numeric keys if table or report schema exists
        # parse item.table_cells_json if present (not used for simple numeric)

        # Build per-unit values and aggregated metrics for this item
        # For generic handling, check submissions payload first
        per_unit = {}
        # union keys: numeric_value, payload 'values' keys, narrative
        # We'll produce one aggregated block representing numeric summary if numeric values exist
        # Numeric aggregation
        numeric_values = {}
        # collect all keys present in payload 'values'
        keys_present = set()
        for key, subs in subs_index[item_id].items():
            # pick last submission by submitted_at as default per-unit value
            last = max(subs, key=lambda x: x.submitted_at or datetime.min)
            # attempt numeric: use numeric_value if set, else try payload json
            num = None
            if last.numeric_value is not None:
                num = Decimal(str(last.numeric_value))
            else:
                try:
                    payload = last.payload_json and json.loads(last.payload_json) or {}
                    values = payload.get('values') if isinstance(payload.get('values'), dict) else {}
                    # if values contains single numeric entry, treat it
                    if values:
                        for k, v in values.items():
                            try:
                                dv = Decimal(str(v))
                                numeric_values.setdefault(k, {})[key] = dv
                                keys_present.add(k)
                            except Exception:
                                pass
                except Exception:
                    pass
            if num is not None:
                numeric_values.setdefault('value', {})[key] = num

        # construct aggregated output
        fields_agg = []
        # handle numeric scalar 'value'
        if 'value' in numeric_values:
            vals_map = numeric_values['value']
            nums = [v for v in vals_map.values() if v is not None]
            agg = {
                'sum': float(sum(nums)) if nums else None,
                'avg': float(sum(nums) / len(nums)) if nums else None,
                'count': len(nums),
                'coverage_pct': (len(nums) / max(1, len(expected_units) or len(expected_assignments))) * 100 if (expected_units or expected_assignments) else None,
            }
            fields_agg.append({'key': 'value', 'type': 'number', 'per_unit': [{ 'unit_key': k, 'value': float(v) } for k, v in vals_map.items()], 'aggregated': agg})

        # handle payload keys
        for k in sorted(keys_present):
            map_k = numeric_values.get(k, {})
            nums = [v for v in map_k.values() if v is not None]
            agg = {
                'sum': float(sum(nums)) if nums else None,
                'avg': float(sum(nums) / len(nums)) if nums else None,
                'count': len(nums),
                'coverage_pct': (len(nums) / max(1, len(expected_units) or len(expected_assignments))) * 100 if (expected_units or expected_assignments) else None,
            }
            per_unit_list = [{ 'unit_key': u, 'value': float(map_k[u]) } for u in map_k.keys()]
            fields_agg.append({'key': k, 'type': 'number', 'per_unit': per_unit_list, 'aggregated': agg})

        # narrative preview: last narrative per unit
        narrative_per_unit = []
        for key, subs in subs_index[item_id].items():
            last = max(subs, key=lambda x: x.submitted_at or datetime.min)
            nav = last.narrative_content or ''
            narrative_per_unit.append({'unit_key': key, 'narrative': nav, 'submitted_at': last.submitted_at})

        items_out.append({'item_id': item_id, 'item_code': item.item_code, 'title': item.title, 'fields': fields_agg, 'narratives': narrative_per_unit})

    # build units list
    units_out = []
    if expected_units:
        for k, v in expected_units.items():
            units_out.append({'unit_key': k, 'unit_id': v.get('unit_id'), 'unit_name': v.get('label')})
    else:
        for k, v in expected_assignments.items():
            units_out.append({'unit_key': k, 'assignment_id': v.get('assignment_id'), 'label': v.get('label')})

    return {'task_id': task_id, 'cycle_key': cycle_key, 'items': items_out, 'units': units_out}



def export_outline_docx(task_id, cycle_key=None):
    """Export a minimal DOCX with per-item per-unit narrative + numeric aggregated table.

    Returns (bytes, filename) tuple.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt

    agg = build_aggregate_context(task_id, cycle_key=cycle_key)
    if 'error' in agg:
        raise ValueError('task_not_found')

    # simple document
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    task_title = f"Task {task_id}"
    doc.add_heading(task_title, level=1)

    for item in agg['items']:
        doc.add_heading(f"{item.get('item_code') or ''} {item.get('title')}", level=2)
        # narratives
        if item.get('narratives'):
            for n in item['narratives']:
                doc.add_paragraph(f"{n.get('unit_key')}: { (n.get('narrative') or '').strip() }")
        # numeric fields aggregated
        for f in item.get('fields', []):
            doc.add_paragraph(f"Field: {f.get('key')}")
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text = 'Unit'
            hdr[1].text = 'Value'
            for pu in f.get('per_unit', []):
                row = table.add_row().cells
                row[0].text = str(pu.get('unit_key'))
                row[1].text = str(pu.get('value'))
            # aggregated
            agg_block = f.get('aggregated') or {}
            doc.add_paragraph('Aggregated: ' + ', '.join(f"{k}: {v}" for k, v in agg_block.items() if v is not None))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"task_{task_id}_outline.docx"
    return bio.read(), filename
