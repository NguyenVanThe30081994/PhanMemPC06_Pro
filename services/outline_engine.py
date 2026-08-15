# -*- coding: utf-8 -*-
"""
Engine phân tích đề cương (outline): tiêu đề, phân cấp, nhận diện người nhận,
trường số liệu trong đề cương.

Tách từ routes/tasks.py (Pha 2). routes/tasks.py vẫn re-export toàn bộ tên cũ
nên các chỗ gọi hiện có không đổi.
"""

import io
import json
import os
import re
import zlib

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import fitz as PdfDocument  # pymupdf
    if not hasattr(PdfDocument, "open"):
        PdfDocument = None
except ImportError:
    PdfDocument = None

from models import AppRole, User
from services.task_categories import _task_assignment_unit_options
from services.task_scope import _parse_bulk_child_task_titles
from utils import remove_accents

TASK_OUTLINE_ALLOWED_EXTENSIONS = {".docx", ".txt", ".pdf"}


def _clean_outline_title(raw_value):
    cleaned = str(raw_value or "").strip()
    if not cleaned:
        return ""
    cleaned = re.sub(r"^\s*(?:[-*+•]\s*|\+\s*|(?:[0-9]{1,3}\.){1,4}\s*|[0-9]{1,3}[.)]\s*|[A-Za-z][.)]\s*|[IVXLCDMivxlcdm]+[.)]\s*)", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .:-")
    return cleaned

def _is_outline_structural_heading(raw_text, cleaned_text):
    raw_text = str(raw_text or "").strip()
    cleaned_text = str(cleaned_text or "").strip()
    if not cleaned_text:
        return True

    normalized = remove_accents(cleaned_text.replace("Đ", "D").replace("đ", "d")).lower()
    compact = re.sub(r"\s+", " ", normalized).strip()
    has_bullet_prefix = bool(re.match(r"^\s*(?:[-*+•]|\+)\s*", raw_text))
    has_multi_numbering = bool(re.match(r"^\s*(?:(?:[0-9]{1,3}\.){1,4}|[IVXLCDMivxlcdm]+\.)", raw_text))

    if compact.startswith("de cuong bao cao") or compact.startswith("trong trien khai"):
        return True

    if re.match(r"^\s*[IVXLCDMivxlcdm]+\.\s+", raw_text):
        return True

    if re.match(r"^\s*[0-9]{1,3}\.\s+[A-ZĂÂĐÊÔƠƯÀÁẢÃẠẮẰẲẴẶẤẦẨẪẬÈÉẺẼẸẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌỐỒỔỖỘỚỜỞỠỢÙÚỦŨỤỨỪỬỮỰỲÝỶỸỴ\s]+$", raw_text):
        return True

    structural_markers = (
        "nhan xet, danh gia",
        "ve hoan thien the che",
        "ve cai cach tthc",
        "ve phat trien kinh te xa hoi",
        "ve phat trien cong dan so",
        "ve ket noi, chia se, tao lap du lieu",
        "ve nguon nhan luc",
        "ve trien khai cac mo hinh diem cua de an 06",
        "ve du lieu",
        "ve ha tang cong nghe thong tin",
        "ve an ninh an toan",
        "ve kinh phi",
        "ve nguon nhan luc",
        "trien khai cac giai phap thanh toan khong dung tien mat",
        "trien khai cac cong cu so va tien ich so cho nguoi dan",
        "pho cap ky nang so",
        "co che khuyen khich cong dan tham gia tren moi truong so",
        "trung tam phuc vu hanh chinh cong",
        "thue tinh",
    )
    if compact in structural_markers:
        return True

    if compact.startswith("cac so, ban, nganh") and compact.endswith("bao cao ket qua") and len(compact) < 60 and ":" not in compact:
        return True
    if compact.startswith("cac so, ban, nganh") and compact.endswith("bao cao ve") and len(compact) < 60 and ":" not in compact:
        return True
    # Cảnh giác: dòng nội dung cũng có thể bắt đầu bằng "Các sở, ban, ngành,
    # Ủy ban nhân dân xã, phường báo cáo..." nhưng LÀ NỘI DUNG cần gán (vd mục
    # 7.2, 8 trong đề cương). Chỉ coi là tiêu đề mục khi dòng ngắn, kiểu tiêu đề
    # (kết thúc bằng "báo cáo kết quả" / "báo cáo về") và không có dấu hai chấm.
    if compact.startswith("cac so, ban, nganh, uy ban nhan dan xa, phuong") and not has_bullet_prefix:
        if len(compact) < 60 and ":" not in compact and (
            compact.endswith("bao cao ket qua") or compact.endswith("bao cao ve")
        ):
            return True
    if compact.startswith("voi chinh phu") or compact.startswith("voi bo, nganh trung uong") or compact.startswith("voi uy ban nhan dan tinh") or compact.startswith("voi so, ban, nganh"):
        return True

    if has_multi_numbering and len(cleaned_text) < 40 and ":" not in cleaned_text and not any(
        keyword in compact for keyword in ("bao cao", "ton tai", "nhiem vu trong tam", "kien nghi", "de xuat")
    ):
        return True

    return False

def _parse_outline_docx_titles(file_storage):
    if DocxDocument is None:
        raise ValueError("Máy chủ chưa cài thư viện đọc file Word (.docx).")

    try:
        file_storage.stream.seek(0)
        file_bytes = file_storage.stream.read()
        document = DocxDocument(io.BytesIO(file_bytes))
    except Exception:
        raise ValueError("Không đọc được file đề cương Word. Hãy thử lại với file .docx rõ nội dung đầu mục.")

    candidates = []
    for paragraph in document.paragraphs:
        raw_text = str(getattr(paragraph, "text", "") or "").strip()
        if not raw_text:
            continue
        cleaned = _clean_outline_title(raw_text)
        if len(cleaned) < 3:
            continue
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
        is_outline_like = bool(
            re.match(r"^\s*(?:[-*+•]|\+|(?:[0-9]{1,3}\.){1,4}|[0-9]{1,3}[.)]|[A-Za-z][.)]|[IVXLCDMivxlcdm]+[.)])\s*", raw_text)
            or any(token in style_name for token in ("heading", "list", "bullet", "number"))
        )
        if is_outline_like and not _is_outline_structural_heading(raw_text, cleaned):
            candidates.append(cleaned)

    if not candidates:
        for paragraph in document.paragraphs:
            raw_text = str(getattr(paragraph, "text", "") or "").strip()
            if not raw_text:
                continue
            cleaned = _clean_outline_title(raw_text)
            if len(cleaned) < 3 or _is_outline_structural_heading(raw_text, cleaned):
                continue
            candidates.append(cleaned)

    return _parse_bulk_child_task_titles("\n".join(candidates))

def _pdf_decode_string_token(token):
    """Giải mã chuỗi văn bản PDF (nội dung giữa cặp ngoặc) — xử lý escape \\(, \\), \\\\, octal."""
    out = bytearray()
    i = 0
    n = len(token)
    while i < n:
        ch = token[i]
        if ch == 0x5C:  # backslash
            if i + 1 >= n:
                break
            nxt = token[i + 1]
            simple = {0x6E: 0x0A, 0x72: 0x0D, 0x74: 0x09, 0x62: 0x08, 0x66: 0x0C}  # n r t b f
            if nxt in simple:
                out.append(simple[nxt])
                i += 2
            elif nxt in (0x28, 0x29, 0x5C):  # ( ) \
                out.append(nxt)
                i += 2
            elif 0x30 <= nxt <= 0x37:  # octal \ddd
                j = i + 1
                octal = 0
                count = 0
                while j < n and count < 3 and 0x30 <= token[j] <= 0x37:
                    octal = octal * 8 + (token[j] - 0x30)
                    j += 1
                    count += 1
                out.append(octal & 0xFF)
                i = j
            else:
                out.append(nxt)
                i += 2
        else:
            out.append(ch)
            i += 1
    return bytes(out)


def _pdf_text_stdlib(data):
    """Trích văn bản từ PDF bằng thư viện chuẩn (zlib) — fallback khi máy chủ chưa
    cài pymupdf. Chỉ xử lý được PDF có text stream FlateDecode (không phải ảnh chụp)."""
    import zlib

    page_texts = []
    for m in re.finditer(rb"stream\r?\n", data):
        start = m.end()
        end_marker = data.find(b"endstream", start)
        if end_marker < 0:
            continue
        raw = data[start:end_marker]
        header = data[max(0, m.start() - 300):m.start()]
        if b"/FlateDecode" not in header:
            continue
        content = None
        for candidate in (raw, raw.rstrip(b"\r\n\x00 ")):
            try:
                content = zlib.decompress(candidate)
                break
            except Exception:
                continue
        if not content:
            continue
        # Chỉ xử lý content stream thật (có BT/ET đánh dấu bắt đầu văn bản).
        # Tránh chạy regex trên stream nhị phân (ảnh/font) chứa byte "Tj"/"TJ" ngẫu nhiên.
        if b"BT" not in content or not (b"Tj" in content or b"TJ" in content):
            continue
        text_parts = []
        for tm in re.finditer(rb"\((?:\\.|[^()])*\)\s*Tj|\[(?:[^\[\]]*)\]\s*TJ", content):
            token = tm.group(0)
            if token.endswith(b"Tj"):
                inner = token[token.find(b"(") + 1:token.rfind(b")")]
                text_parts.append(_pdf_decode_string_token(inner))
            else:
                arr_inner = token[1:token.rfind(b"]")]
                for sm in re.finditer(rb"\((?:\\.|[^()])*\)", arr_inner):
                    inner = sm.group(0)[1:-1]
                    text_parts.append(_pdf_decode_string_token(inner))
        if text_parts:
            page_texts.append(b"".join(text_parts).decode("utf-8", errors="replace"))
    return "\n".join(page_texts)


def _parse_outline_pdf_text(file_storage):
    """Trích dòng chữ từ file PDF: ưu tiên pymupdf, fallback thư viện chuẩn.
    Trả về (lines, error) — lines rỗng kèm error nếu không đọc được."""
    try:
        file_storage.stream.seek(0)
        data = file_storage.stream.read()
    except Exception:
        return [], "Không đọc được file PDF."
    if not data:
        return [], "File PDF rỗng."
    lines = []
    if PdfDocument is not None:
        try:
            document = PdfDocument.open(stream=data, filetype="pdf")
            try:
                for page in document:
                    for line in (str(getattr(page, "get_text", lambda: "")() or "")).splitlines():
                        cleaned = str(line or "").strip()
                        if cleaned:
                            lines.append(cleaned)
            finally:
                try:
                    document.close()
                except Exception:
                    pass
        except Exception:
            lines = []
    if not lines:
        raw_text = _pdf_text_stdlib(data)
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        if PdfDocument is None:
            return [], (
                "Máy chủ chưa cài thư viện đọc PDF (pymupdf). Hãy chạy: pip install pymupdf "
                "trên máy chủ rồi thử lại."
            )
        return [], (
            "File PDF không có nội dung chữ để phân tích (có thể là file ảnh chụp/scanned). "
            "Hãy tải bản .docx hoặc file PDF có chữ rõ ràng."
        )
    return lines, None


def _parse_outline_pdf_titles(file_storage):
    lines, error = _parse_outline_pdf_text(file_storage)
    if error:
        raise ValueError(error)
    return _parse_bulk_child_task_titles("\n".join(lines))


def _parse_outline_text_titles(file_storage):
    try:
        file_storage.stream.seek(0)
        raw_bytes = file_storage.stream.read()
    except Exception:
        raise ValueError("Không đọc được file đề cương văn bản.")

    raw_text = ""
    for encoding in ("utf-8", "utf-8-sig", "cp1258"):
        try:
            raw_text = raw_bytes.decode(encoding)
            break
        except Exception:
            raw_text = ""
    if not raw_text:
        raise ValueError("File đề cương văn bản không đúng định dạng UTF-8.")
    return _parse_bulk_child_task_titles(raw_text)

def _parse_outline_upload_titles(file_storage):
    if not file_storage or not getattr(file_storage, "filename", ""):
        return []

    extension = os.path.splitext(file_storage.filename or "")[1].lower()
    if extension not in TASK_OUTLINE_ALLOWED_EXTENSIONS:
        raise ValueError("Chỉ hỗ trợ đề cương dạng .docx, .txt hoặc .pdf.")

    if extension == ".docx":
        return _parse_outline_docx_titles(file_storage)
    if extension == ".pdf":
        return _parse_outline_pdf_titles(file_storage)
    return _parse_outline_text_titles(file_storage)

OUTLINE_ASSIGNEE_HINT_KEYWORDS = (
    "đơn vị thực hiện",
    "cơ quan thực hiện",
    "đơn vị chủ trì",
    "đơn vị",
    "giao cho",
    "người thực hiện",
    "cán bộ phụ trách",
    "người phụ trách",
    "phụ trách",
    "chủ trì",
    "thực hiện",
    "phối hợp",
    "bộ phận",
)
OUTLINE_ASSIGNEE_NORM_KEYWORDS = (
    "don vi thuc hien",
    "co quan thuc hien",
    "don vi chu tri",
    "don vi",
    "giao cho",
    "nguoi thuc hien",
    "can bo phu trach",
    "nguoi phu trach",
    "phu trach",
    "chu tri",
    "thuc hien",
    "phoi hop",
    "bo phan",
)

def _normalize_outline_match_text(value):
    text = str(value or "").replace("Đ", "D").replace("đ", "d")
    normalized = remove_accents(text)
    normalized = normalized.lower()
    normalized = re.sub("[.,;:()\\[\\]\"'“”‘’]", " ", normalized)
    normalized = re.sub(r"[\s\-_/|]+", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()

def _task_assignment_catalog():
    """Danh mục đơn vị / vai trò / cán bộ để đối sánh 'giao cho ai' trong đề cương."""
    catalog = {"units": [], "roles": [], "users": []}
    seen_keys = set()
    for item in _task_assignment_unit_options():
        key = (item.get("value") or item.get("stable_value") or item.get("name") or "").strip()
        name = (item.get("name") or item.get("value") or key or "").strip()
        if not key or key in seen_keys or len(_normalize_outline_match_text(name)) < 3:
            continue
        seen_keys.add(key)
        catalog["units"].append({"key": key, "name": name, "match": _normalize_outline_match_text(name)})
    for role in AppRole.query.order_by(AppRole.name.asc()).all():
        role_name = str(role.name or "").strip()
        if len(_normalize_outline_match_text(role_name)) < 3:
            continue
        catalog["roles"].append({"id": role.id, "name": role_name, "match": _normalize_outline_match_text(role_name)})
    for user in User.query.filter(User.is_active.is_(True)).order_by(User.fullname.asc()).all():
        matches = []
        for label in (user.fullname, user.username):
            normalized = _normalize_outline_match_text(label)
            if len(normalized) >= 3 and normalized not in matches:
                matches.append(normalized)
        if matches:
            catalog["users"].append({"id": user.id, "fullname": user.fullname, "matches": matches})
    return catalog

def _find_all_outline_assignee_matches(normalized_text, catalog):
    """Tìm mọi đơn vị / vai trò / cá nhân xuất hiện trong chuỗi, không trùng lặp."""
    candidates = []
    for user in catalog["users"]:
        for label in user["matches"]:
            if len(label) >= 3:
                candidates.append(("user", label, user))
    for role in catalog["roles"]:
        if len(role["match"]) >= 3:
            candidates.append(("role", role["match"], role))
    for unit in catalog["units"]:
        if len(unit["match"]) >= 3:
            candidates.append(("unit", unit["match"], unit))
    candidates.sort(key=lambda item: len(item[1]), reverse=True)

    found = []
    for kind, label, target in candidates:
        search_from = 0
        while True:
            idx = normalized_text.find(label, search_from)
            if idx == -1:
                break
            if not any(idx < end and idx + len(label) > start for start, end, _kind, _target in found):
                found.append((idx, idx + len(label), kind, target))
                break
            search_from = idx + 1
    found.sort(key=lambda item: item[0])
    return found

def _resolve_outline_assignee_hint(hint_text, catalog):
    """Nhận diện cấu hình gán việc (đơn vị / vai trò / cá nhân) từ một đoạn chữ."""
    if not hint_text or not str(hint_text).strip():
        return None

    raw = str(hint_text).strip()
    keyword_pattern = (
        r"(?:đơn vị thực hiện|cơ quan thực hiện|đơn vị chủ trì|giao cho|người thực hiện|"
        r"cán bộ phụ trách|người phụ trách|đơn vị|phụ trách|chủ trì|thực hiện|phối hợp|bộ phận)"
        r"\s*[:：]\s*(.+)"
    )
    keyword_match = re.search(keyword_pattern, raw, re.IGNORECASE)
    if keyword_match:
        raw = keyword_match.group(1)

    normalized = _normalize_outline_match_text(raw)
    matched = {"units": [], "roles": [], "users": []}
    matched_labels = []
    for _start, _end, kind, target in _find_all_outline_assignee_matches(normalized, catalog):
        if kind == "unit":
            if target["key"] not in matched["units"]:
                matched["units"].append(target["key"])
                matched_labels.append(target["name"])
        elif kind == "role":
            if target["id"] not in matched["roles"]:
                matched["roles"].append(target["id"])
                matched_labels.append(target["name"])
        else:
            if target["id"] not in matched["users"]:
                matched["users"].append(target["id"])
                matched_labels.append(target["fullname"])

    if matched["users"]:
        return {"assign_type": "user", "unit_domains": [], "role_ids": [], "user_ids": matched["users"], "labels": matched_labels}
    if matched["roles"]:
        return {"assign_type": "role", "unit_domains": [], "role_ids": matched["roles"], "user_ids": [], "labels": matched_labels}
    if matched["units"]:
        return {"assign_type": "unit", "unit_domains": matched["units"], "role_ids": [], "user_ids": [], "labels": matched_labels}
    return None

def _strip_outline_assignee_suffix(title, catalog):
    """Tách phần 'giao cho ai' nằm ngay trong tiêu đề đầu mục (nếu có)."""
    raw = str(title or "").strip()
    for separator in (" — ", " – ", " - ", " (", "(", "[", ": "):
        if separator not in raw:
            continue
        left, right = raw.split(separator, 1)
        right = right.strip().rstrip(")]")
        if _resolve_outline_assignee_hint(right, catalog):
            return left.strip(), right
    return raw, ""

def _looks_like_outline_assignee_text(text, catalog):
    normalized = _normalize_outline_match_text(text)
    if len(normalized) < 3:
        return False
    if any(keyword in normalized for keyword in OUTLINE_ASSIGNEE_NORM_KEYWORDS):
        return True
    return _resolve_outline_assignee_hint(text, catalog) is not None

def _resolve_outline_rows_assignments(rows, catalog):
    resolved = []
    seen = set()
    for row in rows:
        title = str(row.get("title") or "").strip()
        if not title:
            continue
        dedupe_key = _normalize_outline_match_text(title)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        hint = " | ".join(part for part in [row.get("hint") or ""] if part)
        assignment = _resolve_outline_assignee_hint(hint, catalog)
        resolved.append(
            {
                "title": title[:255],
                "assign_type": assignment["assign_type"] if assignment else "",
                "domain": "",
                "unit_domains": assignment["unit_domains"] if assignment else [],
                "role_ids": assignment["role_ids"] if assignment else [],
                "user_ids": assignment["user_ids"] if assignment else [],
                "assignee_hint": hint,
                "assignee_detected": bool(
                    assignment
                    and (assignment["unit_domains"] or assignment["role_ids"] or assignment["user_ids"])
                ),
            }
        )
    return resolved

def _paragraph_is_outline_item(text):
    raw_text = str(text or "").strip()
    if not raw_text:
        return False
    cleaned = _clean_outline_title(raw_text)
    if len(cleaned) < 3:
        return False
    if re.match(r"^\s*[IVXLCDMivxlcdm]+\.\s+", raw_text):
        return False
    if re.match(r"^\s*(?:[-*+•]|\+)\s*", raw_text):
        return True
    if re.match(r"^\s*(?:(?:[0-9]{1,3}\.){1,4}|[0-9]{1,3}[.)]|[A-Za-z][.)])\s+", raw_text):
        return True
    return False


def _is_outline_heading(text, style_name=""):
    raw = str(text or "").strip()
    if not raw:
        return False
    # DOCX heading style
    if style_name and any(token in str(style_name).lower() for token in ("heading", "title", "đề mục", "tieu de")):
        return True
    # Common heading patterns: A. I. 1. 1.1. a) (1)
    if re.match(r"^\s*(?:[A-Z][\.\)])\s*\S", raw):
        return True
    if re.match(r"^\s*(?:[IVXLCDM]+[\.\)])\s*\S", raw):
        return True
    if re.match(r"^\s*(?:\d{1,3}\.){1,4}\s+\S", raw):
        return True
    if re.match(r"^\s*\d{1,3}[\.\)]\s+\S", raw):
        return True
    if re.match(r"^\s*[a-z][\.\)]\s+\S", raw, re.IGNORECASE):
        return True
    if re.match(r"^\s*\(\d{1,3}\)\s+\S", raw):
        return True


def _get_heading_level(raw_text):
    """Return the hierarchy level and normalized marker for a heading."""
    raw = str(raw_text or "").strip()
    if not raw:
        return (0, "")
    # Roman/letter chapters are top-level containers. Chấp nhận dấu chấm tùy
    # chọn (vd: "III KIẾN NGHỊ, ĐỀ XUẤT") vì nhiều đề cương viết thiếu dấu chấm.
    roman_match = re.match(r"^\s*([IVXLCDM]+)\.?\s+", raw)
    if roman_match:
        return (1, roman_match.group(1))
    letter_match = re.match(r"^\s*([A-Z])\.\s+", raw)
    if letter_match:
        return (1, letter_match.group(1))
    # Numeric headings: 1., 1), 1.1, 1.1., 1.1.1, ...
    # Numeric levels start at 2 so they nest below Roman/letter chapters.
    numeric_match = re.match(r"^\s*(\d{1,3}(?:\.\d{1,3})*)[\.\)]?\s+(.+)$", raw)
    if numeric_match:
        number_part = numeric_match.group(1)
        return (number_part.count(".") + 2, number_part)
    return (0, "")

def _parse_outline_with_hierarchy(paragraphs, is_docx=False):
    """Parse đề cương với hierarchy đa cấp.
    
    Cấu trúc:
    - Level 1: I, II, III hoặc A, B, C (chương/phần lớn)
    - Level 2: 1, 2, 3 (mục lớn)
    - Level 3: 1.1, 1.2, 1.1.1 (mục con)
    - Content: các gạch đầu dòng (+, -, •) thuộc về mục cha gần nhất
    
    Trả về danh sách items với structure:
    [
        {
            "level": 1,
            "title": "I. KẾT QUẢ CÁC MẶT CÔNG TÁC",
            "number": "I",
            "children": [
                {
                    "level": 2,
                    "title": "1. CÔNG TÁC THAM MƯU...",
                    "number": "1",
                    "children": [
                        {
                            "level": 3,
                            "title": "1.1. Các Sở, ban, ngành...",
                            "number": "1.1",
                            "content_lines": ["- Dòng 1", "+ Dòng 2"],  # các gạch đầu dòng
                            "children": []
                        }
                    ]
                }
            ]
        }
    ]
    """
    items = []
    stack = []  # Stack để track parent items: [(level, item), ...]
    
    for para in paragraphs:
        if is_docx:
            text = str(getattr(para, "text", "") or "").strip()
            style_name = str(getattr(getattr(para, "style", None), "name", "") or "").strip()
        else:
            text = str(para or "").strip()
            style_name = ""
        
        if not text:
            continue
        
        # Check if this is a heading
        level, number = _get_heading_level(text)
        
        if level > 0:
            # Đây là heading
            title = text.strip()
            # Clean the title by removing the number prefix for storage
            clean_title = re.sub(r"^\s*(?:[IVXLCDM]+|[A-Z]|\d{1,3}(?:\.\d{1,3})*)[\.\)]?\s*", "", title).strip()
            
            new_item = {
                "level": level,
                "title": clean_title[:255],
                "full_title": title[:255],
                "number": number,
                "content_lines": [],
                "bullets": [],
                "children": []
            }
            
            # Pop stack until we find parent with level < current level
            while stack and stack[-1][0] >= level:
                stack.pop()
            
            # Add to parent's children or root
            if stack:
                parent_item = stack[-1][1]
                parent_item["children"].append(new_item)
            else:
                items.append(new_item)
            
            stack.append((level, new_item))
        else:
            # Đây là content line (gạch đầu dòng). Giữ nguyên cấp lồng nhau:
            # - / • / * : gạch đầu dòng cấp 1 (1 nội dung để gán)
            # +         : nội dung con nằm trong gạch đầu dòng cấp 1 liền trước
            if stack:
                current_parent = stack[-1][1]
                if not _is_outline_structural_heading(text, text):
                    _append_outline_bullet(current_parent, text)
            elif items:
                _append_outline_bullet(items[0], text)
    
    return items


def _append_outline_bullet(item, raw_text):
    """Thêm một dòng nội dung vào mục, giữ cấu trúc cha/con:
    gạch đầu dòng '-' là bullet cấp 1; dòng '+' là con của bullet cấp 1 liền trước.
    Dòng thường (không có ký hiệu) cũng là bullet cấp 1.
    """
    text = str(raw_text or "").strip()
    if not text:
        return
    item.setdefault("content_lines", []).append(text)
    bullets = item.setdefault("bullets", [])
    if re.match(r"^\s*\+", text):
        if bullets:
            bullets[-1].setdefault("children", []).append({"text": text, "type": "plus", "children": []})
        else:
            bullets.append({"text": text, "type": "plus", "children": []})
    else:
        bullet_type = "dash" if re.match(r"^\s*[-–—•*]", text) else "para"
        bullets.append({"text": text, "type": bullet_type, "children": []})


def _flatten_hierarchy_to_rows(hierarchy_items, catalog=None):
    """Chuyển cây đề cương thành danh sách dòng để gán việc.

    Mỗi GẠCH ĐẦU DÒNG cấp 1 ('-', '•', '*') là MỘT nội dung để gán (row riêng).
    Dòng '+' nằm dưới một gạch đầu dòng là NỘI DUNG CON của gạch đó (row con,
    có parent_row_index trỏ về row cha). Mặc định gán cho gạch đầu dòng sẽ tự
    gán cho các nội dung con, nhưng quản trị vẫn sửa được riêng từng dòng con.

    - Tiêu đề row = chính nội dung gạch đầu dòng (đã bỏ ký hiệu), không cần lặp
      lại tiêu đề mục vì đường dẫn mục đã đủ chi tiết (vd: I. » 1. » 1.1. ...).
    - Mục lá không có gạch đầu dòng -> tự nó là 1 việc (đầu mục chỉ có tiêu đề).
    - Mục trung gian (chỉ chứa mục con) -> không tạo việc, chỉ đệ quy xuống mục con.
    """
    rows = []
    seen = set()

    def make_row(text, full_heading, number, level, parent_row_index=None):
        raw = str(text or "").strip()
        if not raw:
            return None
        # Bỏ ký hiệu gạch đầu dòng ở đầu dòng: '-', '–', '—', '•', '*', '+'
        title = re.sub(r"^\s*(?:[-–—•*+]\s*|\+\s*)\s*", "", raw).strip()
        title = re.sub(r"\s+", " ", title).strip(" .:")
        if len(title) < 3:
            return None
        dedupe = _normalize_outline_match_text(f"{full_heading} {title}")
        if dedupe in seen:
            return None
        seen.add(dedupe)
        full_text = f"{full_heading}\n{raw}" if raw else full_heading
        assignment = _resolve_outline_assignee_hint(full_text, catalog) if catalog else None
        number_fields = _extract_number_fields_from_text(raw)
        return {
            "title": title[:255],
            "content": raw[:3000],
            "heading": full_heading[:255],
            "level": level,
            "number": number,
            "parent_row_index": parent_row_index,
            "has_numbers": bool(number_fields),
            "number_fields": number_fields,
            "skeleton": _outline_skeleton_text(raw[:3000], number_fields),
            "assign_type": assignment["assign_type"] if assignment else "",
            "domain": "",
            "unit_domains": assignment["unit_domains"] if assignment else [],
            "role_ids": assignment["role_ids"] if assignment else [],
            "user_ids": assignment["user_ids"] if assignment else [],
            "assignee_hint": full_text[:500],
            "assignee_detected": bool(assignment and (assignment["unit_domains"] or assignment["role_ids"] or assignment["user_ids"])),
        }

    def process_item(item, parent_heading=""):
        # Build full heading path
        full_heading = item.get("full_title", item.get("title", ""))
        if parent_heading:
            full_heading = f"{parent_heading} » {full_heading}"

        cleaned_title = _clean_outline_title(item.get("title", ""))
        number = str(item.get("number") or "").strip()
        bullets = item.get("bullets") or []
        has_children = bool(item.get("children"))

        if bullets:
            # Mỗi gạch đầu dòng cấp 1 là 1 nội dung để gán; '+' là nội dung con.
            for bullet in bullets:
                parent_row = make_row(bullet.get("text"), full_heading, number, item.get("level", 3))
                parent_row_index = None
                if parent_row:
                    parent_row_index = len(rows)
                    rows.append(parent_row)
                for child in bullet.get("children") or []:
                    child_row = make_row(
                        child.get("text"),
                        full_heading,
                        number,
                        item.get("level", 3) + 1,
                        parent_row_index=parent_row_index,
                    )
                    if child_row:
                        rows.append(child_row)
        elif not has_children and len(cleaned_title) >= 3:
            # Mục lá không có gạch đầu dòng -> tự nó là 1 việc (đầu mục chỉ có tiêu đề).
            row_title = f"{number}. {cleaned_title}" if number else cleaned_title
            row = make_row(row_title, full_heading, number, item.get("level", 1))
            if row:
                rows.append(row)
        # Process children
        for child in item.get("children", []):
            process_item(child, full_heading)

    for item in hierarchy_items:
        process_item(item)

    return rows


    return False


_OUTLINE_METRIC_KEYWORDS = [
    "tổng số", "tổng", "số lượng", "số", "đạt", "có", "trên", "dưới", "vượt", "chiếm",
    "tỷ lệ", "tỉ lệ", "tỷ suất", "tỉ suất", "phần trăm", "%", "bằng", "đến", "trong đó",
    "lũy kế", "còn", "đã", "được", "giải quyết", "tiếp nhận", "xử lý", "hoàn thành",
    "tăng", "giảm", "so với", "mức", "chỉ số", "kpi", "chỉ tiêu", "dư nợ", "người",
    "hồ sơ", "lượt", "trạm", "tài khoản", "đơn vị", "cơ sở", "yêu cầu", "thông tin",
    "trường hợp", "khoản", "đồng", "tỷ", "triệu", "căn cước", "thẻ", "tài khoản",
]

_OUTLINE_NUMBER_TOKEN = r"\d{1,3}(?:\.\d{3})*(?:,\d+)?|\d+(?:,\d+)?"

_OUTLINE_UNIT_STOPWORDS = {
    "và", "của", "trong", "đến", "so", "với", "theo", "đạt", "chiếm", "từ", "đã",
    "được", "có", "tổng", "số", "là", "các", "khoảng", "gồm", "năm", "tháng", "ngày",
    "trên", "dưới", "vượt", "bằng", "tăng", "giảm", "còn", "để", "không", "tại", "về",
    "đồng", "trong đó", "toàn", "tỉnh", "huyện", "xã", "phường", "cấp", "kỳ", "thời điểm",
}


def _mask_outline_dates_and_years(text):
    """Thay ngày tháng, năm và số hiệu văn bản bằng ký tự cùng độ dài để
    không trùng với số liệu báo cáo. Phân số thật (54.105/57.417) không bị che."""
    masked = list(text)
    for match in re.finditer(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b", text):
        for i in range(match.start(), match.end()):
            masked[i] = "#"
    for match in re.finditer(r"(?<![\d/.])(?:19|20)\d{2}(?![\d/.])", text):
        for i in range(match.start(), match.end()):
            masked[i] = "#"
    # Số hiệu văn bản: 66.7/2025, 18/2023, 05/2025/NQ-CP... (RHS 2-4 chữ số, không phải phân số)
    for match in re.finditer(r"\d{1,3}(?:\.\d{1,3})*/\d{2,4}(?![\d.,])", text):
        for i in range(match.start(), match.end()):
            masked[i] = "#"
    # Số hiệu văn bản dạng 7709/QĐ-CAT-ANM, 18/CT-TTg: số + "/" theo sau bởi chữ cái
    for match in re.finditer(r"\d{2,}(?:\.\d+)*/(?=[A-Za-zÀ-Ỹà-ỹ])", text):
        for i in range(match.start(), match.end()):
            masked[i] = "#"
    return "".join(masked)


def _outline_number_metric(text, start, end, value):
    """Đánh giá 1 số/1 cặp số có phải số liệu báo cáo (metric) không."""
    if value.endswith("%") or "/" in value:
        return True
    compact = value.replace(".", "").replace(",", "").replace("%", "")
    if len(compact) >= 6:
        return True
    before = text[max(0, start - 25):start].lower()
    has_keyword = any(keyword in before for keyword in _OUTLINE_METRIC_KEYWORDS)
    unit = _outline_number_unit(text, end)
    if len(compact) >= 3 and (has_keyword or unit):
        return True
    if len(compact) == 2 and has_keyword and unit:
        return True
    return False


def _outline_number_unit(text, end):
    """Lấy đơn vị theo sau số (vd: %, tỷ đồng, người, hồ sơ)."""
    after = text[end:end + 40]
    match = re.match(r"\s*(%|%%)\s*", after)
    if match:
        return "%"
    match = re.match(r"\s*([\w\u00C0-\u1EF9]+(?:\s+[\w\u00C0-\u1EF9]+)?)", after)
    if not match:
        return ""
    unit = match.group(1).strip()
    first_word = unit.split()[0].lower()
    if first_word in _OUTLINE_UNIT_STOPWORDS or re.match(r"^(năm|tháng|ngày)$", first_word):
        return ""
    if unit.endswith(",") or unit.endswith(";") or unit.endswith("."):
        unit = unit[:-1]
    return unit[:30]


def _extract_number_fields_from_text(text):
    """Trích xuất các trường số liệu từ nội dung đề cương.

    Trả về danh sách dict: {blank_id, label, value, unit, kind, start, end}.
    - Cặp X/Y (54.105/57.417) -> 1 ô trống, kind="pair", value="X/Y".
    - Ngày tháng (13/7/2026), năm (2026), số hiệu văn bản (18/CT-TTg, số 66.7/2025)
      bị loại.
    - start/end là khoảng vị trí trong text gốc để thay ô trống / merge lại.
    """
    if not text:
        return []
    masked = _mask_outline_dates_and_years(text)
    fields = []
    seen_spans = set()
    blank_id = 0
    number_pattern = re.compile(_OUTLINE_NUMBER_TOKEN)
    pair_pattern = re.compile(
        r"(" + _OUTLINE_NUMBER_TOKEN + r")\s*/\s*(" + _OUTLINE_NUMBER_TOKEN + r")"
    )
    index = 0
    while index < len(text):
        pair = pair_pattern.match(masked, index)
        if pair:
            side1 = pair.group(1).replace(".", "").replace(",", "")
            side2 = pair.group(2).replace(".", "").replace(",", "")
            if re.match(r"^(19|20)\d{2}$", side1) or re.match(r"^(19|20)\d{2}$", side2):
                # Cặp chứa năm (số hiệu văn bản 66.7/2025, 18/2023...) — bỏ qua
                index = pair.end()
                continue
            value = pair.group(1) + "/" + pair.group(2)
            start, end = pair.span()
            if _outline_number_metric(text, start, end, value):
                blank_id += 1
                fields.append(
                    _outline_build_number_field(text, start, end, value, "pair", blank_id)
                )
                seen_spans.add((start, end))
                index = end
                continue
        number = number_pattern.match(masked, index)
        if not number:
            index += 1
            continue
        value = number.group(0)
        start, end = number.span()
        if (start, end) in seen_spans:
            index = end
            continue
        # Token khớp thiếu: số dài hơn bị cắt (vd "7709" -> "770") — bỏ qua để tránh sai lệch
        if end < len(masked) and masked[end].isdigit():
            index = end
            continue
        # Số hiệu văn bản dạng 18/CT-TTg, 66.7/2025 — theo sau là "/" (có thể có khoảng trắng)
        if re.match(r"\s*/", masked[end:]):
            index = end
            continue
        compact = value.replace(".", "").replace(",", "").replace("%", "")
        if len(compact) < 2 and not value.endswith("%"):
            index = end
            continue
        kind = "percent" if value.endswith("%") else "plain"
        if _outline_number_metric(text, start, end, value):
            blank_id += 1
            fields.append(
                _outline_build_number_field(text, start, end, value, kind, blank_id)
            )
            seen_spans.add((start, end))
        index = end
    return fields


def _outline_build_number_field(text, start, end, value, kind, blank_id):
    before = text[max(0, start - 80):start].strip()
    label = before.split(";")[-1].split(".")[-1].split(",")[-1].strip()
    label = re.sub(
        r"^\s*(?:và|hoặc|cùng|các|của|để|được|đã|có|là|từ|trong|đến)\s+",
        "",
        label,
        flags=re.IGNORECASE,
    ).strip()
    if len(label) < 2 or len(label) > 120:
        label = before[-60:].strip() if len(before) > 60 else before.strip()
        label = re.sub(
            r"^\s*(?:và|hoặc|cùng|các|của|để|được|đã|có|là|từ|trong|đến)\s+",
            "",
            label,
            flags=re.IGNORECASE,
        ).strip()
    label = re.sub(r"\s+(?:là|của|và|đạt|có|các)$", "", label, flags=re.IGNORECASE).strip()
    if len(label) < 2:
        label = f"Số liệu {blank_id}"
    unit = _outline_number_unit(text, end)
    return {
        "blank_id": blank_id,
        "label": label[:120],
        "value": value,
        "unit": unit[:30],
        "kind": kind,
        "start": start,
        "end": end,
    }


def _parse_vn_number(text):
    """Parse số theo cả định dạng VN (1.234,5 / 85,5) lẫn quốc tế (85.5 / 1234.5)."""
    if text is None:
        return None
    text = str(text).strip().replace("%", "").replace(" ", "")
    if not text or not re.match(r"^[\d.,]+$", text):
        return None
    try:
        if "," in text and "." in text:
            return float(text.replace(".", "").replace(",", "."))
        if "," in text:
            return float(text.replace(",", "."))
        if "." in text:
            parts = text.split(".")
            if len(parts) > 2 or (len(parts) == 2 and len(parts[1]) == 3 and len(parts[0]) > 1):
                return float(text.replace(".", ""))
            return float(text)
        return float(text)
    except ValueError:
        return None


def _parse_outline_blank_value(text):
    """Giá trị ô trống: chuỗi thô nếu là cặp X/Y, float nếu là số thường; None nếu lỗi."""
    if not text:
        return None
    text = str(text).strip()
    pair = re.match(r"^([\d.,]+)\s*/\s*([\d.,]+)$", text)
    if pair and _parse_vn_number(pair.group(1)) is not None and _parse_vn_number(pair.group(2)) is not None:
        return text
    return _parse_vn_number(text)


def _outline_blank_numeric(value):
    """Giá trị số của 1 ô trống để cộng gộp (cặp X/Y lấy tử số)."""
    if value is None:
        return None
    text = str(value).strip()
    match = re.match(r"([\d.,]+)", text)
    if not match:
        return None
    return _parse_vn_number(match.group(1))


def _outline_sources_json(sources):
    if not sources:
        return None
    try:
        return json.dumps([str(source).strip() for source in sources if str(source).strip()], ensure_ascii=False)
    except Exception:
        return None


def _outline_skeleton_text(text, fields):
    """Văn bản với mỗi số liệu thay bằng dấu [...] để xem trước trong wizard."""
    if not fields:
        return text
    result = []
    cursor = 0
    for field in sorted(fields, key=lambda f: f.get("start", 0)):
        start = int(field.get("start", 0))
        end = int(field.get("end", 0))
        if start < cursor or start > len(text) or end > len(text):
            continue
        result.append(text[cursor:start])
        result.append("[...]")
        cursor = end
    result.append(text[cursor:])
    return "".join(result)


