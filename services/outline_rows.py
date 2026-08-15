# -*- coding: utf-8 -*-
"""
Phân tích đề cương thành dòng (rows): chia block, nhận diện bảng, xuất row.

Tách từ routes/tasks.py (Pha 2). routes/tasks.py vẫn re-export toàn bộ tên cũ
nên các chỗ gọi hiện có không đổi.
"""

import io
import os
import re

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import fitz as PdfDocument  # pymupdf
    if not hasattr(PdfDocument, "open"):
        PdfDocument = None
except ImportError:
    PdfDocument = None

from services.outline_engine import (
    TASK_OUTLINE_ALLOWED_EXTENSIONS,
    _clean_outline_title,
    _extract_number_fields_from_text,
    _flatten_hierarchy_to_rows,
    _is_outline_heading,
    _looks_like_outline_assignee_text,
    _normalize_outline_match_text,
    _outline_skeleton_text,
    _parse_outline_pdf_text,
    _parse_outline_with_hierarchy,
    _resolve_outline_assignee_hint,
    _strip_outline_assignee_suffix,
    _task_assignment_catalog,
)
from utils import remove_accents


def _split_outline_paragraphs_into_blocks(paragraphs, is_docx=False):
    """Chia danh sách paragraphs thành các block (heading + content paragraphs)."""
    blocks = []
    current_heading = None
    current_content = []
    current_heading_text = ""
    current_heading_style = ""

    for para in paragraphs:
        if is_docx:
            text = str(getattr(para, "text", "") or "").strip()
            style_name = str(getattr(getattr(para, "style", None), "name", "") or "").strip()
        else:
            text = str(para or "").strip()
            style_name = ""
        if not text:
            continue
        if _is_outline_heading(text, style_name):
            if current_heading is not None or current_content:
                blocks.append({
                    "heading": current_heading_text,
                    "content": " ".join(current_content).strip(),
                })
            current_heading = text
            current_heading_text = text
            current_content = []
        else:
            if current_heading is None:
                # Leading text before any heading: treat as heading with no title
                current_heading = ""
                current_heading_text = ""
            current_content.append(text)

    if current_heading is not None or current_content:
        blocks.append({
            "heading": current_heading_text,
            "content": " ".join(current_content).strip(),
        })
    return blocks

OUTLINE_TABLE_ROLE_LABELS = {
    "stt": "Số thứ tự",
    "content": "Nội dung nhiệm vụ",
    "lead": "Đơn vị chủ trì",
    "coordinate": "Đơn vị phối hợp",
    "deadline": "Thời gian",
    "product": "Sản phẩm, kết quả",
    "note": "Ghi chú",
    "other": "Cột khác",
}


def _table_build_schema(header_cells):
    """Dựng cấu trúc cột bảng từ dòng tiêu đề: mỗi cột có index/header/role/visible.

    - role: tự nhận diện qua _table_column_role (content/lead/coordinate/deadline/...)
    - visible: cột có hiển thị cho đơn vị nhận hay không (mặc định hiện content/lead/
      coordinate/deadline; cột Stt, Sản phẩm, Ghi chú, cột khác mặc định ẩn — quản trị
      có thể tích/bỏ tích trong wizard).
    """
    roles = _table_column_role(header_cells)
    schema = []
    for idx, header in enumerate(header_cells):
        role = next((role for role, col_idx in roles.items() if col_idx == idx), "other")
        if role == "index":
            role = "stt"
        visible = role in ("content", "lead", "coordinate", "deadline")
        schema.append(
            {
                "index": idx,
                "header": re.sub(r"\s+", " ", str(header or "").strip())[:200],
                "role": role,
                "visible": visible,
            }
        )
    return schema


def _table_column_role(cells):
    """Dò vai trò từng cột của bảng theo dòng tiêu đề (bảng không có cột Stt).
    Trả về {vai_trò: chỉ_số_cột} (vd: {"content": 0, "lead": 1, "deadline": 2})."""
    roles = {}
    for idx, header in enumerate(cells):
        key = remove_accents(str(header or "").strip().lower())
        key = re.sub(r"[^a-z0-9 ]", " ", key)
        key = re.sub(r"\s+", " ", key).strip()
        if key in ("stt", "tt", "so", "so thu tu"):
            roles.setdefault("index", idx)
        elif key == "noi dung" or "noi dung" in key or "nhiem vu" in key or "cong viec" in key or key == "viec":
            roles.setdefault("content", idx)
        elif "chu tri" in key or key in ("don vi", "on vi") or ("thuc hien" in key and ("don vi" in key or "on vi" in key)):
            roles.setdefault("lead", idx)
        elif "phoi hop" in key:
            roles.setdefault("coordinate", idx)
        elif "thoi gian" in key or "thoi han" in key or "thoi diem" in key:
            roles.setdefault("deadline", idx)
        elif "san pham" in key or "ket qua" in key:
            roles.setdefault("product", idx)
        elif "ghi chu" in key:
            roles.setdefault("note", idx)
    return roles


def _table_header_based_rows(table, catalog=None, seen=None):
    """Xử lý bảng KHÔNG có cột số thứ tự (Stt/La Mã): dò vai trò cột theo tiêu đề
    (vd: Nhiệm vụ | Đơn vị | Thời hạn) và biến mỗi dòng dữ liệu thành 1 nội dung gán.
    """
    seen = seen if seen is not None else set()
    rows = []
    data_rows = list(table.rows)
    if not data_rows:
        return rows
    first_cells = [re.sub(r"\s+", " ", (c.text or "").strip().replace("\n", " ")) for c in data_rows[0].cells]
    roles = _table_column_role(first_cells)
    schema = _table_build_schema(first_cells)
    start = 1 if roles else 0
    for row in data_rows[start:]:
        cells = [re.sub(r"\s+", " ", (c.text or "").strip().replace("\n", " ")) for c in row.cells]
        if not any(cells):
            continue
        if roles:
            content = cells[roles["content"]] if roles.get("content", -1) >= 0 else ""
            lead = cells[roles["lead"]] if roles.get("lead", -1) >= 0 else ""
            coordinate = cells[roles["coordinate"]] if roles.get("coordinate", -1) >= 0 else ""
            deadline = cells[roles["deadline"]] if roles.get("deadline", -1) >= 0 else ""
            product = cells[roles["product"]] if roles.get("product", -1) >= 0 else ""
            if not content:
                # Không tìm thấy cột nội dung rõ ràng -> lấy ô có nội dung dài nhất
                content = max(cells, key=len) if cells else ""
        else:
            content = max(cells, key=len) if cells else ""
            lead = coordinate = deadline = product = ""
        if not content:
            continue
        unit_domains = []
        if catalog and lead:
            assignment = _resolve_outline_assignee_hint(f"Cơ quan chủ trì: {lead}", catalog)
            if assignment:
                unit_domains = assignment.get("unit_domains") or []
        content_parts = [content]
        if lead:
            content_parts.append(f"Cơ quan chủ trì: {lead}")
        if coordinate:
            content_parts.append(f"Cơ quan phối hợp: {coordinate}")
        if deadline:
            content_parts.append(f"Thời gian: {deadline}")
        if product:
            content_parts.append(f"Sản phẩm, kết quả: {product}")
        raw = f"- {' | '.join(content_parts)}"
        dedupe = _normalize_outline_match_text(f" {content}")
        if dedupe in seen:
            continue
        seen.add(dedupe)
        rows.append(
            {
                "title": content[:255],
                "content": raw[:3000],
                "heading": "",
                "level": 2,
                "number": "",
                "parent_row_index": None,
                "has_numbers": False,
                "number_fields": [],
                "assign_type": "unit" if unit_domains else "",
                "domain": "",
                "unit_domains": unit_domains,
                "role_ids": [],
                "user_ids": [],
                "assignee_hint": f"Cơ quan chủ trì: {lead}" if lead else "",
                "assignee_detected": bool(unit_domains),
                "table_schema": schema,
                "table_cells": {str(idx): cell for idx, cell in enumerate(cells)},
            }
        )
    return rows


def _table_rows_to_outline_rows(document, catalog=None):
    """Chuyển các BẢNG nhiệm vụ trong đề cương (cột: Stt | Nội dung nhiệm vụ |
    Cơ quan, đơn vị chủ trì | Cơ quan, đơn vị phối hợp | Thời gian | Sản phẩm, kết quả | Ghi chú)
    thành các dòng gán việc:
    - Dòng mục (ô đầu là số La Mã I, II, III...) -> tiêu đề mục (heading).
    - Dòng nhiệm vụ (ô đầu là số 1, 2, 3...) -> 1 nội dung để gán, ĐƠN VỊ CHỦ TRÌ
      được gán sẵn từ cột "Cơ quan, đơn vị chủ trì" (khớp với danh mục đơn vị).
    - Bảng không có cột Stt -> dò vai trò cột theo tiêu đề (fallback).
    """
    rows = []
    seen = set()
    for table in document.tables:
        current_heading = ""
        table_had_numeric = False
        header_cells = [re.sub(r"\s+", " ", (c.text or "").strip().replace("\n", " ")) for c in table.rows[0].cells] if table.rows else []
        schema = _table_build_schema(header_cells)
        roles = _table_column_role(header_cells)
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", (c.text or "").strip().replace("\n", " ")) for c in row.cells]
            if not cells or not cells[0]:
                continue
            first = cells[0].strip()
            if first.lower().startswith("stt") or first.lower() == "tt":
                continue
            # Dòng mục: ô đầu là số La Mã
            if re.match(r"^[IVXLCDM]+$", first):
                title = next((c for c in cells[1:] if c and c.strip() and c.strip() != first), "")
                if title:
                    current_heading = f"{first}. {title.strip()}"[:255]
                continue
            # Dòng nhiệm vụ: ô đầu là số thứ tự
            if re.match(r"^\d{1,3}$", first):
                content_index = roles.get("content")
                lead_index = roles.get("lead")
                coordinate_index = roles.get("coordinate")
                deadline_index = roles.get("deadline")
                product_index = roles.get("product")
                if content_index is None:
                    content_index = 1 if len(cells) > 1 else 0
                content = cells[content_index].strip() if content_index < len(cells) else ""
                lead = cells[lead_index].strip() if lead_index is not None and lead_index < len(cells) else ""
                coordinate = cells[coordinate_index].strip() if coordinate_index is not None and coordinate_index < len(cells) else ""
                deadline = cells[deadline_index].strip() if deadline_index is not None and deadline_index < len(cells) else ""
                product = cells[product_index].strip() if product_index is not None and product_index < len(cells) else ""
                if not content:
                    continue
                # Gán sẵn đơn vị chủ trì (chỉ cột Chủ trì, không lấy cột Phối hợp)
                unit_domains = []
                if catalog and lead:
                    assignment = _resolve_outline_assignee_hint(f"Cơ quan chủ trì: {lead}", catalog)
                    if assignment:
                        unit_domains = assignment.get("unit_domains") or []
                content_parts = [content]
                if lead:
                    content_parts.append(f"Cơ quan chủ trì: {lead}")
                if coordinate:
                    content_parts.append(f"Cơ quan phối hợp: {coordinate}")
                if deadline:
                    content_parts.append(f"Thời gian: {deadline}")
                if product:
                    content_parts.append(f"Sản phẩm, kết quả: {product}")
                raw = f"- {' | '.join(content_parts)}"
                dedupe = _normalize_outline_match_text(f"{current_heading} {content}")
                if dedupe in seen:
                    continue
                seen.add(dedupe)
                table_had_numeric = True
                rows.append(
                    {
                        "title": content[:255],
                        "content": raw[:3000],
                        "heading": current_heading[:255],
                        "level": 2,
                        "number": first,
                        "parent_row_index": None,
                        "has_numbers": False,
                        "number_fields": [],
                        "assign_type": "unit" if unit_domains else "",
                        "domain": "",
                        "unit_domains": unit_domains,
                        "role_ids": [],
                        "user_ids": [],
                        "assignee_hint": f"Cơ quan chủ trì: {lead}",
                        "assignee_detected": bool(unit_domains),
                        "table_schema": schema,
                        "table_cells": {str(idx): cell for idx, cell in enumerate(cells)},
                    }
                )
        # Bảng không có dòng số thứ tự nào -> thử dò cột theo tiêu đề
        if not table_had_numeric:
            rows.extend(_table_header_based_rows(table, catalog=catalog, seen=seen))
    return rows


def _parse_outline_docx_rows(file_storage):
    """Parse Word docx với hierarchy awareness (đoạn văn + bảng nhiệm vụ)."""
    if DocxDocument is None:
        raise ValueError("Máy chủ chưa cài thư viện đọc file Word (.docx).")

    try:
        file_storage.stream.seek(0)
        file_bytes = file_storage.stream.read()
        document = DocxDocument(io.BytesIO(file_bytes))
    except Exception:
        raise ValueError("Không đọc được file đề cương Word. Hãy thử lại với file .docx rõ nội dung đầu mục.")

    catalog = _task_assignment_catalog()
    paragraphs = list(document.paragraphs)
    hierarchy_items = _parse_outline_with_hierarchy(paragraphs, is_docx=True)
    rows = _flatten_hierarchy_to_rows(hierarchy_items, catalog=catalog)
    # Gộp thêm các dòng từ BẢNG nhiệm vụ (nếu file Word có bảng)
    if getattr(document, "tables", None):
        table_rows = _table_rows_to_outline_rows(document, catalog=catalog)
        rows.extend(table_rows)
    return rows


def _parse_outline_pdf_rows(file_storage):
    """Parse file PDF (báo cáo / đề cương): trích chữ từng trang -> dòng -> cây mục lục."""
    lines, error = _parse_outline_pdf_text(file_storage)
    if error:
        raise ValueError(error)

    catalog = _task_assignment_catalog()
    hierarchy_items = _parse_outline_with_hierarchy(lines, is_docx=False)
    rows = _flatten_hierarchy_to_rows(hierarchy_items, catalog=catalog)
    # Gộp thêm dòng từ các bảng trong PDF nếu có (bảng nhiệm vụ dạng chữ)
    try:
        file_storage.stream.seek(0)
        document = PdfDocument.open(stream=file_storage.stream.read(), filetype="pdf")
        for page in document:
            try:
                for table in (getattr(page, "find_tables", lambda: None)() or {}).tables:
                    data = table.extract()
                    if not data or not data[0] or not any(data[0]):
                        continue
                    headers = [re.sub(r"\s+", " ", str(c or "").strip().replace("\n", " ")) for c in data[0]]
                    roles = _table_column_role(headers)
                    schema = _table_build_schema(headers)
                    seen = {_normalize_outline_match_text(str(r.get("title") or "")) for r in rows}
                    for data_row in data[1:]:
                        cells = [re.sub(r"\s+", " ", str(c or "").strip().replace("\n", " ")) for c in data_row]
                        if not any(cells):
                            continue
                        content = cells[roles["content"]] if roles.get("content", -1) >= 0 else (max(cells, key=len) if cells else "")
                        if not content:
                            continue
                        lead = cells[roles["lead"]] if roles.get("lead", -1) >= 0 else ""
                        deadline = cells[roles["deadline"]] if roles.get("deadline", -1) >= 0 else ""
                        coordinate = cells[roles["coordinate"]] if roles.get("coordinate", -1) >= 0 else ""
                        product = cells[roles["product"]] if roles.get("product", -1) >= 0 else ""
                        if _normalize_outline_match_text(content) in seen:
                            continue
                        seen.add(_normalize_outline_match_text(content))
                        unit_domains = []
                        if catalog and lead:
                            assignment = _resolve_outline_assignee_hint(f"Cơ quan chủ trì: {lead}", catalog)
                            if assignment:
                                unit_domains = assignment.get("unit_domains") or []
                        content_parts = [content]
                        if lead:
                            content_parts.append(f"Cơ quan chủ trì: {lead}")
                        if coordinate:
                            content_parts.append(f"Cơ quan phối hợp: {coordinate}")
                        if deadline:
                            content_parts.append(f"Thời gian: {deadline}")
                        if product:
                            content_parts.append(f"Sản phẩm, kết quả: {product}")
                        rows.append(
                            {
                                "title": content[:255],
                                "content": f"- {' | '.join(content_parts)}"[:3000],
                                "heading": "",
                                "level": 2,
                                "number": "",
                                "parent_row_index": None,
                                "has_numbers": False,
                                "number_fields": [],
                                "assign_type": "unit" if unit_domains else "",
                                "domain": "",
                                "unit_domains": unit_domains,
                                "role_ids": [],
                                "user_ids": [],
                                "assignee_hint": f"Cơ quan chủ trì: {lead}" if lead else "",
                                "assignee_detected": bool(unit_domains),
                                "table_schema": schema,
                                "table_cells": {str(idx): cell for idx, cell in enumerate(cells)},
                            }
                        )
            except Exception:
                continue
        document.close()
    except Exception:
        pass
    return rows


def _blocks_to_outline_rows(blocks):
    catalog = _task_assignment_catalog()
    rows = []
    seen = set()
    for block in blocks:
        heading = str(block.get("heading") or "").strip()
        content = str(block.get("content") or "").strip()
        if not content and not heading:
            continue
        # Use heading as title fallback, but prefer a short title from content first sentence
        title = heading
        if not title:
            sentences = re.split(r"[.!?]\s+", content)
            title = sentences[0].strip() if sentences else content
        cleaned_title = _clean_outline_title(title)
        if len(cleaned_title) < 3:
            continue
        dedupe = _normalize_outline_match_text(cleaned_title)
        if dedupe in seen:
            continue
        seen.add(dedupe)
        # Try to find assignee hint in heading/content or suffix
        full_text = " ".join([part for part in [heading, content] if part])
        title_for_hint, suffix_hint = _strip_outline_assignee_suffix(cleaned_title, catalog)
        hint_parts = [part for part in [suffix_hint] if part]
        if _looks_like_outline_assignee_text(content, catalog):
            hint_parts.append(content)
        number_fields = _extract_number_fields_from_text(content)
        assignment = _resolve_outline_assignee_hint(" | ".join(hint_parts), catalog)
        rows.append(
            {
                "title": cleaned_title[:255],
                "content": content[:2000],
                "heading": heading[:255],
                "has_numbers": bool(number_fields),
                "number_fields": number_fields,
                "skeleton": _outline_skeleton_text(content[:2000], number_fields),
                "assign_type": assignment["assign_type"] if assignment else "",
                "domain": "",
                "unit_domains": assignment["unit_domains"] if assignment else [],
                "role_ids": assignment["role_ids"] if assignment else [],
                "user_ids": assignment["user_ids"] if assignment else [],
                "assignee_hint": " | ".join(hint_parts),
                "assignee_detected": bool(
                    assignment and (assignment["unit_domains"] or assignment["role_ids"] or assignment["user_ids"])
                ),
            }
        )
    return rows

def _parse_outline_text_rows(file_storage):
    """Parse text file với hierarchy awareness."""
    try:
        file_storage.stream.seek(0)
        raw_bytes = file_storage.stream.read()
    except Exception:
        raise ValueError("Không đọc được file đề cương văn bản.")

    raw_text = ""
    for encoding in ("utf-8", "utf-8-sig", "cp1258"):
        try:
            raw_text = raw_bytes.decode(encoding)
            break
        except Exception:
            raw_text = ""
    if not raw_text:
        raise ValueError("File đề cương văn bản không đúng định dạng UTF-8.")

    lines = [line.strip() for line in raw_text.splitlines()]
    # Parse với hierarchy
    hierarchy_items = _parse_outline_with_hierarchy(lines, is_docx=False)
    catalog = _task_assignment_catalog()
    return _flatten_hierarchy_to_rows(hierarchy_items, catalog=catalog)

def _parse_outline_upload_rows(file_storage):
    """Đọc đề cương (.docx/.txt) -> đầu mục kèm người nhận được tự nhận diện."""
    if not file_storage or not getattr(file_storage, "filename", ""):
        return []

    extension = os.path.splitext(file_storage.filename or "")[1].lower()
    if extension not in TASK_OUTLINE_ALLOWED_EXTENSIONS:
        raise ValueError("Chỉ hỗ trợ đề cương dạng .docx, .txt hoặc .pdf.")

    if extension == ".docx":
        return _parse_outline_docx_rows(file_storage)
    if extension == ".pdf":
        return _parse_outline_pdf_rows(file_storage)
    return _parse_outline_text_rows(file_storage)

