# -*- coding: utf-8 -*-
import unittest
import uuid
from datetime import datetime

from app import app
from models import (
    AppRole,
    Task,
    TaskAssignment,
    TaskComment,
    TaskFormField,
    TaskItem,
    TaskParticipant,
    TaskSubmission,
    User,
    db,
)


class TaskOutlineWordExportTests(unittest.TestCase):
    """Kiểm thử chức năng giao việc báo cáo văn bản (OUTLINE):

    - Ma trận tiến độ đầu mục x đơn vị trên màn chi tiết.
    - Xuất file Word tổng hợp nội dung đã nộp.
    - Trả lại bổ sung kèm lý do.
    """

    def setUp(self):
        self.client = app.test_client()
        self.task_id = None
        self.created_user_ids = []

    def tearDown(self):
        with app.app_context():
            if self.task_id:
                TaskComment.query.filter_by(task_id=self.task_id).delete()
                TaskAssignment.query.filter_by(task_id=self.task_id).update(
                    {TaskAssignment.last_submission_id: None}, synchronize_session=False
                )
                TaskSubmission.query.filter_by(task_id=self.task_id).delete()
                TaskAssignment.query.filter_by(task_id=self.task_id).delete()
                TaskItem.query.filter_by(task_id=self.task_id).delete()
                TaskParticipant.query.filter_by(task_id=self.task_id).delete()
                TaskFormField.query.filter_by(task_id=self.task_id).delete()
                Task.query.filter_by(id=self.task_id).delete()
            for user_id in self.created_user_ids:
                User.query.filter_by(id=user_id).delete()
            db.session.commit()

    def _admin_id(self):
        with app.app_context():
            user = (
                User.query.filter_by(username="admin").first()
                or User.query.filter_by(is_active=True).order_by(User.id.asc()).first()
            )
            self.assertIsNotNone(user)
            return user.id

    def _login(self, user_id, is_admin=False):
        with app.app_context():
            user = db.session.get(User, user_id)
            with self.client.session_transaction() as sess:
                sess["uid"] = user.id
                sess["username"] = user.username
                sess["fullname"] = user.fullname
                sess["unit"] = user.unit_area or ""
                sess["unit_area"] = user.unit_area or ""
                sess["unit_area_ref"] = user.unit_area or ""
                sess["unit_key"] = user.unit_key or ""
                sess["role_id"] = user.role_id
                sess["must_change"] = False
                sess["is_admin"] = is_admin
                sess["session_version"] = int(user.session_version or 0)
                sess["csrf_token"] = "task-outline-csrf"
                sess["last_active"] = datetime.now().timestamp()
                sess["login_nonce"] = "task-outline-test"

    def _create_user(self, unit_area, role_id=None):
        with app.app_context():
            user = User(
                username=f"task_user_{uuid.uuid4().hex[:8]}",
                fullname=f"Cán bộ {uuid.uuid4().hex[:6]}",
                role_id=role_id,
                unit_area=unit_area,
                unit_key=unit_area.lower().replace(" ", "-"),
                is_active=True,
            )
            user.set_password("123456")
            db.session.add(user)
            db.session.commit()
            self.created_user_ids.append(user.id)
            return user.id

    def _create_outline_task(self):
        with app.app_context():
            admin_id = self._admin_id()
            task = Task(
                title="Báo cáo tổng hợp kiểm tra Word",
                content="Mô tả công việc kiểm thử.",
                task_mode="OUTLINE",
                author_id=admin_id,
                author_name="Quản trị",
                priority="Cao",
            )
            db.session.add(task)
            db.session.commit()
            self.task_id = task.id

            unit_a_id = self._create_user("Đội A")
            unit_b_id = self._create_user("Đội B")

            item1 = TaskItem(
                task_id=task.id, item_code="1", title="Đầu mục một",
                is_required=True, output_type="OUTLINE", report_kind="narrative", sort_order=1,
            )
            item2 = TaskItem(
                task_id=task.id, item_code="2", title="Đầu mục hai",
                is_required=True, output_type="OUTLINE", report_kind="number", sort_order=2,
            )
            db.session.add_all([item1, item2])
            db.session.flush()

            assignment1 = TaskAssignment(
                task_id=task.id, task_item_id=item1.id, user_id=unit_a_id,
                assignee_type="unit", status="submitted",
            )
            assignment2 = TaskAssignment(
                task_id=task.id, task_item_id=item2.id, user_id=unit_b_id,
                assignee_type="unit", status="assigned",
            )
            db.session.add_all([assignment1, assignment2])
            db.session.commit()

            submission = TaskSubmission(
                task_id=task.id, task_item_id=task.id, assignment_id=assignment1.id,
                submitted_by=unit_a_id, submission_type="narrative", status="submitted",
                narrative_content="Nội dung báo cáo của Đội A.", submitted_at=datetime.now(),
            )
            db.session.add(submission)
            db.session.commit()
            return assignment1.id

    def _create_bare_outline_task(self):
        with app.app_context():
            admin_id = self._admin_id()
            task = Task(
                title="Đề cương giao việc tự động",
                task_mode="OUTLINE",
                author_id=admin_id,
                author_name="Quản trị",
            )
            db.session.add(task)
            db.session.commit()
            self.task_id = task.id
            return task.id

    def test_outline_import_preview_auto_detects_assignee(self):
        import io
        from unittest.mock import patch

        from docx import Document

        task_id = self._create_bare_outline_task()
        self._login(self._admin_id(), is_admin=True)

        document = Document()
        document.add_paragraph("ĐỀ CƯƠNG CÔNG TÁC")
        document.add_paragraph("1. Triển khai công tác PCCC — Đội A")
        document.add_paragraph("2. Tổng hợp số liệu báo cáo")
        document.add_paragraph("Đơn vị thực hiện: Đội B")
        buffer = io.BytesIO()
        document.save(buffer)

        catalog = {
            "units": [
                {"key": "doi-a", "name": "Đội A", "match": "doi a"},
                {"key": "doi-b", "name": "Đội B", "match": "doi b"},
            ],
            "roles": [],
            "users": [],
        }
        with patch("services.outline_rows._task_assignment_catalog", return_value=catalog):
            response = self.client.post(
                f"/tasks/{task_id}/outline/import-preview",
                data={
                    "outline_file": (io.BytesIO(buffer.getvalue()), "de-cuong.docx"),
                    "child_report_kind": "narrative",
                    "child_attachment_required": "0",
                    "csrf_token": "task-outline-csrf",
                },
                content_type="multipart/form-data",
                follow_redirects=True,
            )
        self.assertEqual(response.status_code, 200)
        html = response.get_data(as_text=True)
        self.assertIn("tự nhận diện người nhận cho 2 đầu mục", html)
        self.assertIn("doi-a", html)
        self.assertIn("doi-b", html)
        self.assertIn("unit", html)

    def test_outline_matrix_and_word_export(self):
        assignment1_id = self._create_outline_task()
        self._login(self._admin_id(), is_admin=True)

        detail = self.client.get(f"/tasks/{self.task_id}")
        self.assertEqual(detail.status_code, 200)
        html = detail.get_data(as_text=True)
        self.assertIn("Tiến độ chung", html)
        self.assertIn("Xuất báo cáo Word", html)
        self.assertIn("Đội A", html)
        self.assertIn("Đội B", html)
        self.assertIn("Đầu mục một", html)

        response = self.client.get(f"/tasks/{self.task_id}/export-outline.docx")
        self.assertEqual(response.status_code, 200)
        self.assertIn("wordprocessingml", response.headers.get("Content-Type", ""))
        body = response.data
        self.assertGreater(len(body), 1000)
        self.assertEqual(body[:2], b"PK")

        with app.app_context():
            assignment = db.session.get(TaskAssignment, assignment1_id)
            self.assertEqual(assignment.status, "submitted")

    def test_return_assignment_records_reason(self):
        assignment1_id = self._create_outline_task()
        self._login(self._admin_id(), is_admin=True)

        response = self.client.post(
            f"/tasks/{self.task_id}/assignments/{assignment1_id}/return",
            data={"csrf_token": "task-outline-csrf", "return_reason": "Thiếu số liệu minh chứng"},
        )
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            assignment = db.session.get(TaskAssignment, assignment1_id)
            self.assertEqual(assignment.status, "returned")
            self.assertIsNotNone(assignment.returned_at)
            comment = (
                TaskComment.query.filter_by(task_id=self.task_id)
                .order_by(TaskComment.id.desc())
                .first()
            )
            self.assertIsNotNone(comment)
            self.assertIn("[TRẢ LẠI]", comment.content)
            self.assertIn("Thiếu số liệu minh chứng", comment.content)

    def test_word_export_requires_permission(self):
        self._create_outline_task()
        outsider_id = self._create_user("Đội Ngoài")
        self._login(outsider_id, is_admin=False)

        response = self.client.get(f"/tasks/{self.task_id}/export-outline.docx")
        self.assertEqual(response.status_code, 302)


if __name__ == "__main__":
    unittest.main()
