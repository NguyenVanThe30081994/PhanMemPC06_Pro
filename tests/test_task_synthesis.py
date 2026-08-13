# -*- coding: utf-8 -*-
import io
import unittest
import uuid
from datetime import datetime

from app import app
from models import (
    Task,
    TaskAssignment,
    TaskComment,
    TaskFormField,
    TaskItem,
    TaskParticipant,
    TaskSubmission,
    User,
    db,
)


class TaskSynthesisTests(unittest.TestCase):
    """Kiểm thử bước tổng hợp báo cáo của quản trị (thay cho tự động gộp):

    - Xem dữ liệu từng đơn vị đã nộp cho 1 đầu mục (synthesis-data).
    - Lưu / xóa văn bản tổng hợp (synthesize).
    - Xuất Word dùng văn bản tổng hợp khi có, fallback gộp tự động khi không.
    """

    def setUp(self):
        self.client = app.test_client()
        self.task_id = None
        self.created_user_ids = []

    def tearDown(self):
        with app.app_context():
            if self.task_id:
                TaskComment.query.filter_by(task_id=self.task_id).delete()
                TaskAssignment.query.filter_by(task_id=self.task_id).update(
                    {TaskAssignment.last_submission_id: None}, synchronize_session=False
                )
                TaskSubmission.query.filter_by(task_id=self.task_id).delete()
                TaskAssignment.query.filter_by(task_id=self.task_id).delete()
                TaskItem.query.filter_by(task_id=self.task_id).delete()
                TaskParticipant.query.filter_by(task_id=self.task_id).delete()
                TaskFormField.query.filter_by(task_id=self.task_id).delete()
                Task.query.filter_by(id=self.task_id).delete()
            for user_id in self.created_user_ids:
                User.query.filter_by(id=user_id).delete()
            db.session.commit()

    def _admin_id(self):
        with app.app_context():
            user = (
                User.query.filter_by(username="admin").first()
                or User.query.filter_by(is_active=True).order_by(User.id.asc()).first()
            )
            self.assertIsNotNone(user)
            return user.id

    def _login(self, user_id, is_admin=False):
        with app.app_context():
            user = db.session.get(User, user_id)
            with self.client.session_transaction() as sess:
                sess["uid"] = user.id
                sess["username"] = user.username
                sess["fullname"] = user.fullname
                sess["unit"] = user.unit_area or ""
                sess["unit_area"] = user.unit_area or ""
                sess["unit_area_ref"] = user.unit_area or ""
                sess["unit_key"] = user.unit_key or ""
                sess["role_id"] = user.role_id
                sess["must_change"] = False
                sess["is_admin"] = is_admin
                sess["session_version"] = int(user.session_version or 0)
                sess["csrf_token"] = "task-synthesis-csrf"
                sess["last_active"] = datetime.now().timestamp()
                sess["login_nonce"] = "task-synthesis-test"

    def _create_user(self, unit_area, role_id=None):
        with app.app_context():
            user = User(
                username=f"task_user_{uuid.uuid4().hex[:8]}",
                fullname=f"Cán bộ {uuid.uuid4().hex[:6]}",
                role_id=role_id,
                unit_area=unit_area,
                unit_key=unit_area.lower().replace(" ", "-"),
                is_active=True,
            )
            user.set_password("123456")
            db.session.add(user)
            db.session.commit()
            self.created_user_ids.append(user.id)
            return user.id

    def _create_task_with_two_units(self):
        with app.app_context():
            admin_id = self._admin_id()
            task = Task(
                title="Báo cáo nhiều đơn vị tổng hợp",
                task_mode="OUTLINE",
                author_id=admin_id,
                author_name="Quản trị",
            )
            db.session.add(task)
            db.session.commit()
            self.task_id = task.id

            unit_a_id = self._create_user("Đội A")
            unit_b_id = self._create_user("Đội B")

            item = TaskItem(
                task_id=task.id, item_code="1", title="Đầu mục chung",
                content="Nội dung gốc của đầu mục.", is_required=True,
                output_type="OUTLINE", report_kind="narrative", sort_order=1,
            )
            db.session.add(item)
            db.session.flush()

            assignment_a = TaskAssignment(
                task_id=task.id, task_item_id=item.id, user_id=unit_a_id,
                assignee_type="unit", status="submitted",
            )
            assignment_b = TaskAssignment(
                task_id=task.id, task_item_id=item.id, user_id=unit_b_id,
                assignee_type="unit", status="submitted",
            )
            db.session.add_all([assignment_a, assignment_b])
            db.session.flush()

            db.session.add_all(
                [
                    TaskSubmission(
                        task_id=task.id, task_item_id=item.id, assignment_id=assignment_a.id,
                        submitted_by=unit_a_id, submission_type="narrative", status="submitted",
                        narrative_content="Đoạn văn của Đội A.", submitted_at=datetime.now(),
                    ),
                    TaskSubmission(
                        task_id=task.id, task_item_id=item.id, assignment_id=assignment_b.id,
                        submitted_by=unit_b_id, submission_type="narrative", status="submitted",
                        narrative_content="Đoạn văn của Đội B.", submitted_at=datetime.now(),
                    ),
                ]
            )
            db.session.commit()
            return item.id

    def test_synthesis_data_lists_each_unit_submission(self):
        item_id = self._create_task_with_two_units()
        self._login(self._admin_id(), is_admin=True)

        response = self.client.get(f"/tasks/{self.task_id}/items/{item_id}/synthesis-data")
        self.assertEqual(response.status_code, 200)
        data = response.get_json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["item"]["item_code"], "1")
        self.assertEqual(data["item"]["synthesis"], "")
        self.assertEqual(len(data["submissions"]), 2)
        units = {sub["unit_name"] for sub in data["submissions"]}
        self.assertEqual(units, {"Đội A", "Đội B"})
        narratives = {sub["narrative"] for sub in data["submissions"]}
        self.assertEqual(narratives, {"Đoạn văn của Đội A.", "Đoạn văn của Đội B."})

    def test_save_synthesis_then_export_uses_synthesis(self):
        from docx import Document as DocxDocument

        item_id = self._create_task_with_two_units()
        self._login(self._admin_id(), is_admin=True)

        response = self.client.post(
            f"/tasks/{self.task_id}/items/{item_id}/synthesize",
            data={"synthesis_content": "Văn bản tổng hợp cuối cùng của quản trị.", "csrf_token": "task-synthesis-csrf"},
        )
        self.assertEqual(response.status_code, 302)

        with app.app_context():
            item = db.session.get(TaskItem, item_id)
            self.assertEqual(item.synthesis_content, "Văn bản tổng hợp cuối cùng của quản trị.")
            self.assertIsNotNone(item.synthesis_updated_at)

        export = self.client.get(f"/tasks/{self.task_id}/export-outline.docx")
        self.assertEqual(export.status_code, 200)
        document = DocxDocument(io.BytesIO(export.data))
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn("Văn bản tổng hợp cuối cùng của quản trị.", texts)
        # Không lặp lại phần từng đơn vị khi đã tổng hợp
        self.assertFalse(any("Đoạn văn của Đội A" in text for text in texts))
        self.assertFalse(any("Đoạn văn của Đội B" in text for text in texts))

    def test_clear_synthesis_falls_back_to_auto_merge(self):
        from docx import Document as DocxDocument

        item_id = self._create_task_with_two_units()
        self._login(self._admin_id(), is_admin=True)

        self.client.post(
            f"/tasks/{self.task_id}/items/{item_id}/synthesize",
            data={"synthesis_content": "Văn bản tổng hợp.", "csrf_token": "task-synthesis-csrf"},
        )
        self.client.post(
            f"/tasks/{self.task_id}/items/{item_id}/synthesize",
            data={"synthesis_content": "", "csrf_token": "task-synthesis-csrf"},
        )

        with app.app_context():
            item = db.session.get(TaskItem, item_id)
            self.assertIsNone(item.synthesis_content)

        export = self.client.get(f"/tasks/{self.task_id}/export-outline.docx")
        document = DocxDocument(io.BytesIO(export.data))
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn("Đoạn văn của Đội A.", texts)
        self.assertIn("Đoạn văn của Đội B.", texts)

    def test_synthesis_requires_manager_permission(self):
        item_id = self._create_task_with_two_units()
        outsider_id = self._create_user("Đội Ngoài")
        self._login(outsider_id, is_admin=False)

        data_response = self.client.get(f"/tasks/{self.task_id}/items/{item_id}/synthesis-data")
        self.assertEqual(data_response.status_code, 403)

        post_response = self.client.post(
            f"/tasks/{self.task_id}/items/{item_id}/synthesize",
            data={"synthesis_content": "Không được phép.", "csrf_token": "task-synthesis-csrf"},
        )
        self.assertEqual(post_response.status_code, 302)

        with app.app_context():
            item = db.session.get(TaskItem, item_id)
            self.assertIsNone(item.synthesis_content)


if __name__ == "__main__":
    unittest.main()
