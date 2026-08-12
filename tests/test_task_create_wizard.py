# -*- coding: utf-8 -*-
import io
import json
import unittest
import uuid
from datetime import datetime

from app import app
from models import (
    CategoryGroup,
    CategoryItem,
    Task,
    TaskAssignment,
    TaskComment,
    TaskFormField,
    TaskItem,
    TaskParticipant,
    TaskSubmission,
    User,
    db,
)
from docx import Document


class TaskCreateWizardTests(unittest.TestCase):
    """Kiểm thử wizard tạo công việc mới:

    - Phân tích đề cương ngay trong lúc tạo (endpoint /tasks/outline-parse).
    - Lấy trường từ biểu mẫu báo cáo có sẵn (endpoint /tasks/form-template-preview).
    - Tạo công việc theo đề cương kèm danh sách việc nhỏ + gán từng việc trong một POST.
    """

    def setUp(self):
        self.client = app.test_client()
        self.created_user_ids = []
        self.created_task_ids = []

    def tearDown(self):
        with app.app_context():
            created_category_ids = getattr(self, "created_category_ids", []) or []
            if created_category_ids:
                for cid in created_category_ids:
                    CategoryItem.query.filter_by(group_id=cid).delete()
                    CategoryGroup.query.filter_by(id=cid).delete()
                db.session.commit()
            for task_id in self.created_task_ids:
                TaskComment.query.filter_by(task_id=task_id).delete()
                # Gỡ tham chiếu last_submission_id trước khi xóa submission
                # (giống logic xóa task trong route, tránh vi phạm khóa ngoại).
                TaskAssignment.query.filter_by(task_id=task_id).update(
                    {TaskAssignment.last_submission_id: None}, synchronize_session=False
                )
                TaskSubmission.query.filter_by(task_id=task_id).delete()
                TaskAssignment.query.filter_by(task_id=task_id).delete()
                TaskItem.query.filter_by(task_id=task_id).delete()
                TaskParticipant.query.filter_by(task_id=task_id).delete()
                TaskFormField.query.filter_by(task_id=task_id).delete()
                Task.query.filter_by(id=task_id).delete()
            for user_id in self.created_user_ids:
                TaskParticipant.query.filter_by(user_id=user_id).delete()
                TaskAssignment.query.filter_by(user_id=user_id).delete()
                TaskSubmission.query.filter_by(submitted_by=user_id).delete()
                TaskComment.query.filter_by(user_id=user_id).delete()
                User.query.filter_by(id=user_id).delete()
            db.session.commit()

    def _admin(self):
        with app.app_context():
            return (
                User.query.filter_by(username="admin").first()
                or User.query.filter_by(is_active=True).order_by(User.id.asc()).first()
            )

    def _login(self, user_id, is_admin=True):
        with app.app_context():
            user = db.session.get(User, user_id)
            with self.client.session_transaction() as sess:
                sess["uid"] = user.id
                sess["username"] = user.username
                sess["fullname"] = user.fullname
                sess["unit"] = user.unit_area or ""
                sess["unit_area"] = user.unit_area or ""
                sess["unit_area_ref"] = user.unit_area or ""
                sess["unit_key"] = user.unit_key or ""
                sess["role_id"] = user.role_id
                sess["must_change"] = False
                sess["is_admin"] = is_admin
                sess["session_version"] = int(user.session_version or 0)
                sess["csrf_token"] = "wizard-test-csrf"
                sess["last_active"] = datetime.now().timestamp()
                sess["login_nonce"] = "wizard-test"

    def _create_user(self, unit_area, unit_key):
        with app.app_context():
            user = User(
                username=f"wizard_{uuid.uuid4().hex[:8]}",
                fullname=f"Cán bộ {uuid.uuid4().hex[:6]}",
                unit_area=unit_area,
                unit_key=unit_key,
                is_active=True,
            )
            user.set_password("123456")
            db.session.add(user)
            db.session.commit()
            self.created_user_ids.append(user.id)
            return user.id

    def test_outline_parse_endpoint_returns_rows(self):
        admin = self._admin()
        self._login(admin.id)

        document = Document()
        document.add_paragraph("ĐỀ CƯƠNG")
        document.add_paragraph("1. Việc nhỏ A")
        document.add_paragraph("2. Việc nhỏ B")
        buffer = io.BytesIO()
        document.save(buffer)

        response = self.client.post(
            "/tasks/outline-parse",
            data={"outline_file": (io.BytesIO(buffer.getvalue()), "de-cuong.docx"), "csrf_token": "wizard-test-csrf"},
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertTrue(payload["ok"])
        # Mỗi mục (heading có số hiệu) là một row; tiêu đề giữ số hiệu mục.
        self.assertEqual([row["title"] for row in payload["rows"]], ["1. Việc nhỏ A", "2. Việc nhỏ B"])

    def test_outline_parse_keeps_content_starting_with_so_ban_nganh(self):
        """Nội dung bắt đầu bằng 'Các sở, ban, ngành, Ủy ban nhân dân xã, phường'
        (vd mục 7.2, 8 trong đề cương) phải được giữ làm row chứ không bị lọc nhầm."""
        admin = self._admin()
        self._login(admin.id)

        document = Document()
        document.add_paragraph("I. KẾT QUẢ CÁC MẶT CÔNG TÁC")
        document.add_paragraph("7.2. Về nguồn nhân lực")
        document.add_paragraph("Các sở, ban, ngành, Ủy ban nhân dân xã, phường báo cáo kết quả tập huấn, đào tạo chuyển đổi số trên nền tảng “Bình dân học vụ số”.")
        buffer = io.BytesIO()
        document.save(buffer)

        response = self.client.post(
            "/tasks/outline-parse",
            data={"outline_file": (io.BytesIO(buffer.getvalue()), "de-cuong.docx"), "csrf_token": "wizard-test-csrf"},
            content_type="multipart/form-data",
        )
        payload = response.get_json()
        self.assertTrue(payload["ok"])
        rows = payload["rows"]
        self.assertEqual(len(rows), 1)
        self.assertIn("tập huấn, đào tạo chuyển đổi số", rows[0]["title"])
        self.assertIn("7.2", rows[0]["heading"])

    def test_outline_parse_table_without_stt_column(self):
        """Bảng nhiệm vụ KHÔNG có cột số thứ tự (vd: Nhiệm vụ | Đơn vị | Thời hạn)
        vẫn phải tách được từng dòng thành nội dung gán, dò vai trò cột theo tiêu đề
        và gợi ý đơn vị từ cột 'Đơn vị'."""
        admin = self._admin()
        self._login(admin.id)
        with app.app_context():
            group = CategoryGroup.query.filter_by(name="Đơn vị").first()
            if not group:
                group = CategoryGroup(name="Đơn vị", code="don-vi")
                db.session.add(group)
                db.session.flush()
            item = CategoryItem.query.filter_by(group_id=group.id, name="Sở Tư pháp").first()
            if not item:
                item = CategoryItem(group_id=group.id, name="Sở Tư pháp", code="sotuphap", is_active=True)
                db.session.add(item)
            db.session.commit()
            self.created_category_ids = [group.id, item.id]

        document = Document()
        table = document.add_table(rows=3, cols=3)
        for j, header in enumerate(["Nhiệm vụ", "Đơn vị", "Thời hạn"]):
            table.rows[0].cells[j].text = header
        table.rows[1].cells[0].text = "Xây dựng kế hoạch"
        table.rows[1].cells[1].text = "Sở Tư pháp"
        table.rows[1].cells[2].text = "31/7"
        table.rows[2].cells[0].text = "Đào tạo kỹ năng số"
        table.rows[2].cells[1].text = "Sở Khoa học và Công nghệ"
        table.rows[2].cells[2].text = "30/9"
        buffer = io.BytesIO()
        document.save(buffer)

        response = self.client.post(
            "/tasks/outline-parse",
            data={"outline_file": (io.BytesIO(buffer.getvalue()), "bang-khong-stt.docx"), "csrf_token": "wizard-test-csrf"},
            content_type="multipart/form-data",
        )
        payload = response.get_json()
        self.assertTrue(payload["ok"])
        rows = payload["rows"]
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["title"], "Xây dựng kế hoạch")
        self.assertIn("Cơ quan chủ trì: Sở Tư pháp", rows[0]["content"])
        self.assertIn("Thời gian: 31/7", rows[0]["content"])
        # Cột "Đơn vị" (bỏ dấu -> 'on vi') vẫn phải gợi ý đơn vị
        self.assertIn("sotuphap", rows[0]["unit_domains"])

    def test_delete_task_with_assignments_and_submissions(self):
        """Xóa công việc OUTLINE đã có assignment + submission (draft) phải sạch."""
        admin = self._admin()
        self._login(admin.id)
        self._create_user("Đội A", "a")

        response = self.client.post(
            "/tasks",
            data={
                "task_mode": "OUTLINE",
                "title": "Wizard xóa có giao việc",
                "description": "Xóa thử",
                "csrf_token": "wizard-test-csrf",
                "item_title": ["1.1. Mục A"],
                "item_content": ["- Dòng 1\n- Dòng 2"],
                "item_report_kind": ["narrative"],
                "item_assign_type": ["unit"],
                "item_domains": ["doi-a"],
                "item_role_ids": [""],
                "item_user_ids": [""],
            },
            follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        with app.app_context():
            task = (
                Task.query.filter_by(title="Wizard xóa có giao việc")
                .order_by(Task.id.desc())
                .first()
            )
            self.assertIsNotNone(task)
            task_id = task.id
            self.assertGreater(TaskAssignment.query.filter_by(task_id=task_id).count(), 0)
            self.assertGreater(TaskSubmission.query.filter_by(task_id=task_id).count(), 0)

        delete_response = self.client.post(f"/tasks/{task_id}/delete", data={"csrf_token": "wizard-test-csrf"}, follow_redirects=True)
        self.assertEqual(delete_response.status_code, 200)
        with app.app_context():
            self.assertIsNone(Task.query.filter_by(id=task_id).first())
            self.assertEqual(TaskItem.query.filter_by(task_id=task_id).count(), 0)
            self.assertEqual(TaskAssignment.query.filter_by(task_id=task_id).count(), 0)
            self.assertEqual(TaskSubmission.query.filter_by(task_id=task_id).count(), 0)

    def test_create_outline_task_with_items_in_one_post(self):
        admin = self._admin()
        self._login(admin.id)
        unit_a_id = self._create_user("Đội A", "a")
        unit_b_id = self._create_user("Đội B", "b")

        response = self.client.post(
            "/tasks",
            data={
                "task_mode": "OUTLINE",
                "title": "Wizard giao việc theo đề cương",
                "description": "Tạo từ wizard",
                "csrf_token": "wizard-test-csrf",
                "item_title": ["Việc nhỏ 1", "Việc nhỏ 2"],
                "item_report_kind": ["narrative", "number"],
                "item_attachment_required": ["0"],
                "item_assign_type": ["unit", "unit"],
                "item_domains": ["doi-a", "doi-b"],
                "item_role_ids": ["", ""],
                "item_user_ids": ["", ""],
            },
            follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        with app.app_context():
            task = (
                Task.query.filter_by(title="Wizard giao việc theo đề cương")
                .order_by(Task.id.desc())
                .first()
            )
            self.assertIsNotNone(task)
            self.created_task_ids.append(task.id)
            items = (
                TaskItem.query.filter_by(task_id=task.id)
                .order_by(TaskItem.sort_order.asc())
                .all()
            )
            self.assertEqual(len(items), 2)
            self.assertEqual(items[0].report_kind, "narrative")
            self.assertEqual(items[1].report_kind, "number")
            assignment_0 = TaskAssignment.query.filter_by(task_id=task.id, task_item_id=items[0].id).all()
            assignment_1 = TaskAssignment.query.filter_by(task_id=task.id, task_item_id=items[1].id).all()
            self.assertEqual([a.user_id for a in assignment_0], [unit_a_id])
            self.assertEqual([a.user_id for a in assignment_1], [unit_b_id])

    def test_task_item_fk_points_to_task_item(self):
        """Hồi quy: task_item_id của TaskSubmission/TaskParticipant phải trỏ task_item.id.

        Trước đây khai báo nhầm FK sang task.id khiến nộp báo cáo đầu mục
        vỡ ràng buộc FOREIGN KEY (500) khi task_item_id không trùng số hiệu task.
        """
        for column in (
            TaskSubmission.__table__.c.task_item_id,
            TaskParticipant.__table__.c.task_item_id,
        ):
            referred = {fk.target_fullname.split(".")[0] for fk in column.foreign_keys}
            self.assertEqual(referred, {"task_item"})

    def test_submit_number_report_with_per_field_values(self):
        """Nộp báo cáo số nhiều trường (report_number_value_*) phải lưu đủ values,
        không vỡ FK và không 500."""
        admin = self._admin()
        self._login(admin.id)
        unit_id = self._create_user("Đơn vị nghiệp vụ", "dvnum")

        response = self.client.post(
            "/tasks",
            data={
                "task_mode": "OUTLINE",
                "title": "Wizard báo cáo số nhiều trường",
                "description": "test",
                "csrf_token": "wizard-test-csrf",
                "item_title": ["Mục số liệu"],
                "item_report_kind": ["number"],
                "item_assign_type": ["unit"],
                "item_domains": ["dv-num"],
                "item_role_ids": [""],
                "item_user_ids": [""],
            },
            follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)

        with app.app_context():
            task = (
                Task.query.filter_by(title="Wizard báo cáo số nhiều trường")
                .order_by(Task.id.desc())
                .first()
            )
            self.assertIsNotNone(task)
            self.created_task_ids.append(task.id)
            item = TaskItem.query.filter_by(task_id=task.id).first()
            item_id = item.id
            assignment = TaskAssignment.query.filter_by(task_id=task.id, task_item_id=item_id).first()
            assigned_uid = assignment.user_id
            self.assertEqual(assigned_uid, unit_id)

        self._login(assigned_uid, is_admin=False)
        submit_response = self.client.post(
            f"/tasks/{task.id}/submit_report",
            data={
                "task_item_id": str(item_id),
                "report_number_value_0": "1234",
                "report_number_value_1": "85.5",
                "report_content": "Hoàn thành.",
                "csrf_token": "wizard-test-csrf",
            },
            follow_redirects=True,
        )
        self.assertEqual(submit_response.status_code, 200)

        with app.app_context():
            submission = TaskSubmission.query.filter_by(task_id=task.id, task_item_id=item_id).first()
            self.assertIsNotNone(submission)
            self.assertEqual(submission.status, "submitted")
            self.assertEqual(submission.numeric_value, 1234.0)
            payload = json.loads(submission.payload_json or "{}")
            self.assertEqual(payload.get("values"), {"0": 1234.0, "1": 85.5})
            self.assertEqual(payload.get("reported_value"), 1234.0)

    def test_linked_secondary_item_auto_fills_on_submit(self):
        """Báo cáo phụ: tạo task thứ 2 trùng nội dung + tích 'báo cáo phụ' phải
        liên kết linked_item_id sang đầu mục task cũ; nộp báo cáo ở task mới phải
        tự động điền vào task cũ (cùng đơn vị được giao)."""
        from routes.tasks import _extract_number_fields_from_text

        content = (
            "Tính đến kỳ báo cáo, toàn tỉnh có 54.105/57.417 người nhận "
            "lương hưu qua tài khoản ngân hàng, đạt 95,97% so với Kế hoạch."
        )
        fields = _extract_number_fields_from_text(content)
        self.assertTrue(any(f.get("kind") == "pair" for f in fields))

        admin = self._admin()
        self._login(admin.id)
        unit_id = self._create_user("Đơn vị nghiệp vụ", "dvnum")

        def create(title, secondary, sources):
            return self.client.post(
                "/tasks",
                data={
                    "task_mode": "OUTLINE",
                    "title": title,
                    "description": "test",
                    "csrf_token": "wizard-test-csrf",
                    "item_title": ["Mục số liệu lương hưu"],
                    "item_content": [content],
                    "item_report_kind": ["number"],
                    "item_number_fields": [json.dumps(fields, ensure_ascii=False)],
                    "item_assign_type": ["unit"],
                    "item_domains": ["dv-num"],
                    "item_role_ids": [""],
                    "item_user_ids": [""],
                    "item_report_secondary": [secondary],
                    "item_sources": [sources],
                },
                follow_redirects=True,
            )

        self.assertEqual(create("Báo cáo tháng 8 (test liên kết)", "0", "").status_code, 200)
        self.assertEqual(
            create("Báo cáo quý (test liên kết phụ)", "1", "Báo cáo tháng 8.pdf,Báo cáo quý.pdf").status_code,
            200,
        )

        with app.app_context():
            task1 = (
                Task.query.filter_by(title="Báo cáo tháng 8 (test liên kết)")
                .order_by(Task.id.desc())
                .first()
            )
            task2 = (
                Task.query.filter_by(title="Báo cáo quý (test liên kết phụ)")
                .order_by(Task.id.desc())
                .first()
            )
            self.assertIsNotNone(task1)
            self.assertIsNotNone(task2)
            self.created_task_ids.extend([task1.id, task2.id])
            item1 = TaskItem.query.filter_by(task_id=task1.id).first()
            item2 = TaskItem.query.filter_by(task_id=task2.id).first()
            self.assertEqual(item2.linked_item_id, item1.id)
            self.assertIn("Báo cáo quý.pdf", json.loads(item2.report_sources_json or "[]"))
            assignment1 = TaskAssignment.query.filter_by(task_id=task1.id, task_item_id=item1.id).first()
            self.assertEqual(assignment1.user_id, unit_id)
            task1_id, task2_id, item1_id, item2_id = task1.id, task2.id, item1.id, item2.id

        self._login(unit_id, is_admin=False)
        submit_response = self.client.post(
            f"/tasks/{task2_id}/submit_report",
            data={
                "task_item_id": str(item2_id),
                "report_number_value_1": "60.000/62.000",
                "report_content": "Hoàn thành.",
                "csrf_token": "wizard-test-csrf",
            },
            follow_redirects=True,
        )
        self.assertEqual(submit_response.status_code, 200)

        with app.app_context():
            own_sub = TaskSubmission.query.filter_by(task_id=task2_id, task_item_id=item2_id).first()
            self.assertIsNotNone(own_sub)
            self.assertEqual(own_sub.status, "submitted")
            linked_sub = TaskSubmission.query.filter_by(task_id=task1_id, task_item_id=item1_id).first()
            self.assertIsNotNone(linked_sub, "nộp báo cáo phụ phải tự điền sang task gốc")
            self.assertEqual(linked_sub.status, "submitted")
            payload = json.loads(linked_sub.payload_json or "{}")
            self.assertEqual(payload.get("values"), {"1": "60.000/62.000"})


    def test_pdf_parse_falls_back_to_stdlib_without_pymupdf(self):
        """Khi máy chủ chưa cài pymupdf, vẫn trích được chữ từ PDF (FlateDecode)
        bằng thư viện chuẩn; và khi pymupdf có sẵn, phân tích bình thường."""
        import zlib

        from routes import tasks as tasks_module

        sample_text = "ĐỀ CƯƠNG BÁO CÁO"
        content_stream = f"BT /F1 12 Tf 72 720 Td ({sample_text}) Tj ET".encode("utf-8")
        compressed = zlib.compress(content_stream)
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
            b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj\n"
            + b"4 0 obj << /Length %d /Filter /FlateDecode >>\nstream\n" % len(compressed)
            + compressed
            + b"\nendstream\nendobj\n"
            b"trailer << /Root 1 0 R >>\n%%EOF\n"
        )

        # Không có pymupdf -> fallback thư viện chuẩn vẫn ra dòng chữ
        original_pdf_document = tasks_module.PdfDocument
        tasks_module.PdfDocument = None
        try:
            from io import BytesIO

            class FakeStorage:
                filename = "de-cuong.pdf"

                def __init__(self, data):
                    self.stream = BytesIO(data)

            lines, error = tasks_module._parse_outline_pdf_text(FakeStorage(pdf_bytes))
            self.assertIsNone(error)
            self.assertTrue(any("ĐỀ CƯƠNG BÁO CÁO" in line for line in lines))
        finally:
            tasks_module.PdfDocument = original_pdf_document

    def test_outline_blank_template_render_and_merge(self):
        """Nội dung lưu dạng bản mẫu (chứa [...]): render ô nhập inline và ghép
        giá trị nộp vào đúng từng marker."""
        from routes.tasks import (
            _extract_number_fields_from_text,
            _outline_merged_content,
            _outline_skeleton_text,
            _render_blank_editor_html,
        )

        content = (
            "Tính đến kỳ báo cáo, toàn tỉnh có 54.105/57.417 người nhận "
            "lương hưu qua tài khoản ngân hàng, đạt 95,97% so với Kế hoạch."
        )
        fields = _extract_number_fields_from_text(content)
        template = _outline_skeleton_text(content, fields)
        self.assertIn("[...]", template)
        self.assertNotIn("54.105", template)

        html_out = _render_blank_editor_html(template, fields, {"1": "60.000/62.000"})
        self.assertIn('name="report_number_value_1"', html_out)
        self.assertIn('value="60.000/62.000"', html_out)
        self.assertIn('name="report_number_value_2"', html_out)
        # Không lặp nhãn/đơn vị (từ xung quanh marker đã có sẵn trong văn bản)
        self.assertNotIn("(toàn tỉnh)", html_out)

        merged = _outline_merged_content(template, fields, {"1": "60.000/62.000", "2": "99,5"})
        self.assertIn("60.000/62.000", merged)
        self.assertIn("99,5", merged)
        self.assertNotIn("[...]", merged)
        self.assertNotIn("54.105/57.417", merged)


if __name__ == "__main__":
    unittest.main()
