# -*- coding: utf-8 -*-
"""
routes/outline.py — Blueprint: Trình biên tập đề cương

Chức năng:
    GET  /outline-editor            -> Trang giao diện (yêu cầu đăng nhập)
    POST /api/parse-outline         -> Upload .docx/.txt, parse cây cấu trúc
    POST /api/save-outline          -> Lưu cây đã chỉnh sửa (JSON -> .docx)

Tích hợp vào app.py:
    from routes.outline import outline_bp
    app.register_blueprint(outline_bp)
"""
import io
import json
import os
import re
import uuid

from flask import (Blueprint, current_app, flash, jsonify, render_template,
                   request, session, url_for)
from werkzeug.utils import secure_filename

from outline_parser import parse_docx, parse_text, build_tree, collect_stats

outline_bp = Blueprint('outline_bp', __name__, url_prefix='')

OUTLINE_ALLOWED_EXTENSIONS = {'.docx', '.txt'}


def _require_login():
    """Các endpoint của blueprint yêu cầu đăng nhập."""
    return bool(session.get('uid'))


def _save_uploaded_file(file_storage):
    """Lưu file upload vào thư mục tạm, trả về đường dẫn."""
    ext = os.path.splitext(file_storage.filename or '')[1].lower()
    if ext not in OUTLINE_ALLOWED_EXTENSIONS:
        raise ValueError('Chỉ hỗ trợ file .docx hoặc .txt.')

    tmp_dir = current_app.config.get('TMP_FOLDER') or 'tmp'
    os.makedirs(tmp_dir, exist_ok=True)
    token = uuid.uuid4().hex[:12]
    fname = f"outline_{token}{ext}"
    path = os.path.join(tmp_dir, fname)
    file_storage.save(path)
    return path


@outline_bp.route('/outline-editor')
def editor_page():
    """Trang trình biên tập đề cương."""
    if not _require_login():
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'error': 'Unauthorized'}), 401
        return ('Bạn cần đăng nhập để dùng chức năng này.', 401)
    return render_template('outline_editor.html')


@outline_bp.route('/api/parse-outline', methods=['POST'])
def api_parse_outline():
    """Upload file đề cương -> JSON cây cấu trúc."""
    if not _require_login():
        return jsonify({'error': 'Unauthorized'}), 401

    if 'file' not in request.files:
        return jsonify({'error': 'Không có file nào được upload.'}), 400

    file_storage = request.files['file']
    if not file_storage or not getattr(file_storage, 'filename', ''):
        return jsonify({'error': 'File không hợp lệ.'}), 400

    try:
        path = _save_uploaded_file(file_storage)
        ext = os.path.splitext(file_storage.filename)[1].lower()

        if ext == '.docx':
            tree = parse_docx(path)
        else:
            with open(path, encoding='utf-8', errors='replace') as f:
                tree = parse_text(f.read())

        # Xoá file tạm sau khi parse xong
        try:
            os.remove(path)
        except OSError:
            pass

        tree['filename'] = file_storage.filename
        tree['stats'] = collect_stats(tree)
        return jsonify(tree)

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        current_app.logger.error(f'parse-outline failed: {e}', exc_info=True)
        return jsonify({'error': 'Lỗi hệ thống khi phân tích file. Hãy thử lại hoặc kiểm tra nhật ký server.'}), 500


@outline_bp.route('/api/save-outline', methods=['POST'])
def api_save_outline():
    """Lưu cây đã chỉnh sửa thành file .docx."""
    if not _require_login():
        return jsonify({'error': 'Unauthorized'}), 401

    payload = request.get_json(silent=True) or {}
    tree = payload.get('tree')
    filename = (payload.get('filename') or 'de-cuong.docx').strip()

    if not tree or not isinstance(tree, dict):
        return jsonify({'error': 'Dữ liệu đề cương không hợp lệ.'}), 400

    try:
        docx_bytes = _tree_to_docx(tree)
    except Exception as e:
        current_app.logger.error(f'save-outline failed: {e}', exc_info=True)
        return jsonify({'error': 'Lỗi hệ thống khi tạo file Word. Hãy thử lại hoặc kiểm tra nhật ký server.'}), 500

    # Lưu ra file trong TASK_FOLDER / uploads để tải về
    save_dir = current_app.config.get('UPLOAD_FOLDER') or 'uploads'
    os.makedirs(save_dir, exist_ok=True)
    safe_name = secure_filename(filename) or 'de-cuong.docx'
    if not safe_name.lower().endswith('.docx'):
        safe_name += '.docx'
    out_path = os.path.join(save_dir, f"outline_{uuid.uuid4().hex[:8]}_{safe_name}")

    with open(out_path, 'wb') as f:
        f.write(docx_bytes)

    return jsonify({
        'success': True,
        'message': 'Đã tạo file Word.',
        'download_url': url_for('outline_bp.api_download_outline', fname=os.path.basename(out_path)),
    })


@outline_bp.route('/api/download-outline/<path:fname>')
def api_download_outline(fname):
    """Tải file .docx đã lưu."""
    if not _require_login():
        return ('Unauthorized', 401)

    from flask import send_file
    base = os.path.basename(fname or '')
    if base != fname or not base.startswith('outline_'):
        return ('Not found', 404)
    path = os.path.join(current_app.config.get('UPLOAD_FOLDER') or 'uploads', base)
    if not os.path.isfile(path):
        return ('Not found', 404)
    return send_file(path, as_attachment=True, download_name=base)


# ── Xuất cây -> .docx ────────────────────────────────────────────────────
def _tree_to_docx(tree):
    """Chuyển cây cấu trúc thành file .docx (python-docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError('Máy chủ chưa cài python-docx.')

    doc = Document()

    # Tiêu đề
    title = (tree.get('title') or '').strip()
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(16)

    subtitle = (tree.get('subtitle') or '').strip()
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.bold = True
        r.font.size = Pt(13)

    def walk(nodes, depth=0):
        for n in nodes:
            kind = n.get('type') or 'para'
            label = (n.get('label') or '').strip()
            text = (n.get('text') or '').strip()
            if kind.startswith('h'):
                try:
                    lvl = int(kind[1:])
                except ValueError:
                    lvl = 2
                title_text = (label + '. ' if label else '') + text
                p = doc.add_heading('', level=min(lvl, 4))
                run = p.add_run(title_text)
                run.bold = True
            elif kind == 'bullet':
                doc.add_paragraph(text, style='List Bullet')
            elif kind == 'plus':
                doc.add_paragraph('+ ' + text, style='List Bullet 2')
            else:
                doc.add_paragraph(text)
            walk(n.get('children') or [], depth + 1)

    walk(tree.get('sections') or [])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
#  GIAO VIỆC THEO ĐỀ CƯƠNG — giữ nguyên cấu trúc cây đa tầng
# ═══════════════════════════════════════════════════════════════════════

@outline_bp.route('/outline-giao-viec')
def giao_viec_page():
    """Trang giao việc theo đề cương (tree view + gán cán bộ/vai trò)."""
    if not _require_login():
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'error': 'Unauthorized'}), 401
        return ('Bạn cần đăng nhập để dùng chức năng này.', 401)
    if not _require_task_process_perm():
        return ('Bạn không có quyền giao công việc.', 403)
    return render_template('outline_assign.html')


# ═══════════════════════════════════════════════════════════════════════
#  GIAO VIỆC THEO ĐỀ CƯƠNG — đơn giản: 1 mục = 1 việc
#
#  Sau khi quét đề cương, mỗi MỤC (heading I., 1., 1.1...) là 1 nhiệm vụ
#  gán cho đơn vị/cán bộ. Các dòng nội dung dưới mục được GỘP LẠI thành
#  nội dung của mục đó (không tách từng dòng thành việc riêng).
#  Không bắt buộc file, không tự gán theo vai trò.
# ═══════════════════════════════════════════════════════════════════════

_HEADING_TYPES = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'h8', 'h9'}


def _is_outline_heading(node):
    return (node or {}).get('type') in _HEADING_TYPES


def _collect_muc_content_lines(node, out=None):
    """
    Gộp TOÀN BỘ các dòng nội dung (bullet/plus/para) dưới một mục,
    kể cả các dòng nằm trong mục con (heading con) — vì 1 mục = 1 việc
    báo cáo nội dung, người được gán cần biết toàn bộ nội dung của mục.
    """
    out = out if out is not None else []
    for child in (node or {}).get('children') or []:
        if _is_outline_heading(child):
            # Thêm tiêu đề mục con làm dòng phân cấp, rồi gom nội dung mục con
            sub_label = (child.get('label') or '').strip()
            sub_text = (child.get('text') or '').strip()
            if sub_label or sub_text:
                out.append('▸ ' + ((sub_label + '. ') if sub_label else '') + sub_text)
            _collect_muc_content_lines(child, out)
            continue
        text = (child.get('text') or '').strip()
        if text:
            prefix = ''
            if child.get('type') == 'bullet':
                prefix = '– '
            elif child.get('type') == 'plus':
                prefix = '+ '
            out.append(prefix + text)
        _collect_muc_content_lines(child, out)
    return out


def _require_task_process_perm():
    """Chỉ người có quyền xử lý module Công việc mới được giao việc (đồng bộ wizard chính)."""
    from services.task_permissions import _can_process_task_module
    return _can_process_task_module()


@outline_bp.route('/api/outline-assignees')
def api_outline_assignees():
    """Danh sách cán bộ để gán việc (chỉ gán trực tiếp cho người, không theo vai trò).

    Danh bạ toàn đơn vị là dữ liệu nhạy cảm: chỉ trả cho người có quyền
    xử lý module Công việc (người thực sự dùng màn giao việc).
    """
    if not _require_login():
        return jsonify({'error': 'Unauthorized'}), 401
    if not _require_task_process_perm():
        return jsonify({'error': 'Bạn không có quyền giao công việc.'}), 403

    from models import User

    users = []
    for u in User.query.filter(User.is_active.is_(True)).order_by(User.fullname.asc()).all():
        users.append({
            'id': u.id,
            'name': (u.fullname or u.username or 'Cán bộ #%s' % u.id).strip(),
            'username': (u.username or '').strip(),
        })

    return jsonify({'users': users})


def _resolve_selected_assignees(assign_info):
    """Chuyển thông tin gán việc thành danh sách user_id (chỉ gán cho cán bộ)."""
    result = []
    ids = (assign_info or {}).get('ids') or []
    for uid in ids:
        if str(uid).isdigit():
            result.append({'user_id': int(uid)})
    return result


@outline_bp.route('/api/create-outline-task', methods=['POST'])
def api_create_outline_task():
    """Tạo công việc OUTLINE từ cây đề cương.

    Mỗi MỤC được gán (heading) tạo 1 TaskItem; nội dung là toàn bộ các dòng
    dưới mục gộp lại. Chỉ tạo việc cho những mục đã được gán người thực hiện.
    """
    if not _require_login():
        return jsonify({'error': 'Unauthorized'}), 401
    if not _require_task_process_perm():
        return jsonify({'error': 'Bạn không có quyền giao công việc.'}), 403

    from datetime import datetime
    from models import Task, TaskAssignment, TaskItem, User, db
    from services.task_scope import _store_assignment_scope
    from services.task_runtime_sync import _ensure_task_runtime_bridge
    from routes.email_service import send_task_assignment_emails
    from utils import log_action, push_notif

    payload = request.get_json(silent=True) or {}
    tree = payload.get('tree')
    title = str(payload.get('title') or '').strip()
    deadline_str = str(payload.get('deadline') or '').strip()
    assignments = payload.get('assignments') or {}   # node_id -> {'ids': [...]}

    if not tree or not isinstance(tree, dict):
        return jsonify({'error': 'Cấu trúc đề cương không hợp lệ.'}), 400
    if not title:
        return jsonify({'error': 'Cần nhập tên công việc.'}), 400
    if not tree.get('sections'):
        return jsonify({'error': 'Đề cương không có đầu mục nào để giao việc.'}), 400

    task = Task(
        title=title[:255],
        content=(tree.get('subtitle') or '')[:2000] or None,
        author_id=session.get('uid'),
        author_name=session.get('fullname', 'Quản trị'),
        priority='Trung bình',
        task_type='Công việc theo đề cương',
        initial_status='Chưa tiếp nhận',
        task_mode='OUTLINE',
    )
    if deadline_str:
        try:
            task.deadline = datetime.strptime(deadline_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    db.session.add(task)
    db.session.flush()

    total_items = 0
    total_assigned = 0
    assignee_user_ids = set()

    def create_items(nodes, parent_item_id=None):
        """Duyệt cây; chỉ tạo TaskItem cho mục (heading) đã được gán."""
        nonlocal total_items, total_assigned
        for n in nodes or []:
            # Chỉ xử lý mục (heading). Nội dung bullet/plus/para được gộp
            # vào mục cha, không tạo việc riêng.
            if not _is_outline_heading(n):
                continue

            assign_info = assignments.get(str(n.get('id'))) if assignments else None
            selected = _resolve_selected_assignees(assign_info) if assign_info else []
            if not selected:
                # Mục chưa gán -> bỏ qua; vẫn duyệt tiếp các mục con
                create_items(n.get('children') or [], parent_item_id)
                continue

            label = (n.get('label') or '').strip()
            text = (n.get('text') or '').strip()
            title_item = (label + '. ' if label else '') + text
            content_lines = _collect_muc_content_lines(n)

            total_items += 1
            item = TaskItem(
                task_id=task.id,
                parent_item_id=parent_item_id,
                item_code=str(total_items),
                title=(title_item or '(chưa có nội dung)')[:255],
                content=('\n'.join(content_lines))[:5000] or None,
                is_required=True,
                output_type='OUTLINE',
                report_kind='narrative',
                attachment_required=False,
                sort_order=total_items,
            )
            db.session.add(item)
            db.session.flush()

            for entry in selected:
                db.session.add(TaskAssignment(
                    task_id=task.id,
                    task_item_id=item.id,
                    user_id=entry['user_id'],
                    assignee_type='user',
                    role_id=None,
                    title_snapshot=item.title,
                    status='assigned',
                    is_required=True,
                    assigned_at=datetime.now(),
                ))
                assignee_user_ids.add(entry['user_id'])
                total_assigned += 1

            create_items(n.get('children') or [], parent_item_id=item.id)

    create_items(tree.get('sections'))

    if total_items == 0:
        db.session.rollback()
        return jsonify({'error': 'Chưa gán mục nào. Hãy gán ít nhất một mục cho cán bộ.'}), 400

    # Lưu phạm vi giao việc theo cá nhân — cùng cơ chế với wizard chính,
    # để lọc/kiểm tra phạm vi nhìn thấy việc nhất quán trên toàn hệ thống.
    try:
        _store_assignment_scope(task, 'user', user_ids=sorted(assignee_user_ids))
    except Exception as scope_error:
        current_app.logger.warning(f'create-outline-task: store scope failed: {scope_error}')

    db.session.commit()

    # Đồng bộ cầu nối runtime (TaskParticipant/submission hiện hành) ngay sau
    # khi tạo, thay vì chờ "vá lười" ở lần xem chi tiết đầu tiên.
    try:
        if _ensure_task_runtime_bridge(task):
            db.session.commit()
    except Exception as bridge_error:
        db.session.rollback()
        current_app.logger.warning(f'create-outline-task: runtime bridge failed: {bridge_error}')

    # Thông báo trong ứng dụng cho từng người được giao.
    for uid in sorted(assignee_user_ids):
        push_notif(uid, 'Công việc mới', f'Bạn vừa được giao: {task.title}', f'/tasks/{task.id}')

    log_action(
        session.get('uid'),
        session.get('fullname', ''),
        'Tạo công việc theo đề cương',
        module='Công việc',
        det=f'{task.title} ({total_items} đầu mục, {total_assigned} lượt giao)',
    )

    # Email thông báo giao việc — không làm hỏng request nếu máy thư lỗi.
    email_sent = 0
    try:
        assignee_users = User.query.filter(User.id.in_(sorted(assignee_user_ids))).all() if assignee_user_ids else []
        base_url = request.host_url.rstrip('/')
        email_result = send_task_assignment_emails(assignee_users, task, base_url=base_url)
        email_sent = len(email_result.get('sent') or [])
    except Exception as email_error:
        current_app.logger.error(f'create-outline-task: send emails failed: {email_error}')

    # Xây link tới chi tiết công việc (an toàn kể cả khi blueprint tasks chưa đăng ký)
    task_url = None
    try:
        task_url = url_for('tasks_bp.task_detail', tid=task.id)
    except Exception:
        task_url = '/tasks/%d' % task.id

    return jsonify({
        'success': True,
        'task_id': task.id,
        'task_url': task_url,
        'items_created': total_items,
        'assignments_created': total_assigned,
        'notifications_created': len(assignee_user_ids),
        'emails_sent': email_sent,
        'message': 'Đã tạo công việc với %d đầu mục, %d lượt giao việc.' % (total_items, total_assigned),
    })
